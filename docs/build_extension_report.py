#!/usr/bin/env python3
# Builds a focused report proposing one evaluative view to extend the QSC Atlas.
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/sessions/gracious-kind-wozniak/mnt/qsc-atlas/docs/QSC_Atlas_Extending_the_Analysis.docx"
INK = RGBColor(0x1A,0x1A,0x1A); MUTED = RGBColor(0x5C,0x5C,0x5C)

doc = Document()
n = doc.styles["Normal"]; n.font.name="Georgia"; n.font.size=Pt(10.5); n.font.color.rgb=INK
n.paragraph_format.space_after=Pt(8); n.paragraph_format.line_spacing=1.25
for hid,sz in (("Heading 1",14),("Heading 2",11.5)):
    st=doc.styles[hid]; st.font.name="Arial"; st.font.size=Pt(sz); st.font.bold=True; st.font.color.rgb=INK
    st.paragraph_format.space_before=Pt(13); st.paragraph_format.space_after=Pt(5)
sec=doc.sections[0]; sec.page_width=Cm(21.0); sec.page_height=Cm(29.7)
for m in ("top_margin","bottom_margin","left_margin","right_margin"): setattr(sec,m,Cm(2.4))

def body(t, after=8):
    p=doc.add_paragraph(); p.add_run(t); p.paragraph_format.space_after=Pt(after); return p
def h1(t): doc.add_heading(t, level=1)

# title
p=doc.add_paragraph(); r=p.add_run("Extending the QSC Atlas: from where a country stands to how credible its approach is")
r.font.name="Georgia"; r.font.size=Pt(17); r.bold=True; p.paragraph_format.space_after=Pt(2)
p=doc.add_paragraph(); r=p.add_run("A proposal for one evaluative view, grounded in the assessment literature")
r.font.name="Arial"; r.font.size=Pt(11.5); r.font.color.rgb=MUTED; p.paragraph_format.space_after=Pt(2)
p=doc.add_paragraph(); r=p.add_run("Working note, June 2026. Companion to the QSC Atlas and to the QSA methods synthesis. [Author and affiliation]")
r.font.name="Arial"; r.font.size=Pt(9); r.font.color.rgb=MUTED; p.paragraph_format.space_after=Pt(12)

ab=doc.add_paragraph(); abr=ab.add_run("Summary. "); abr.bold=True
ab.add_run("The Atlas maps where each country stands in the quantum-safe transition: its coordination bloc, "
  "its standards role, and the instruments that bind it. It does not yet say how good that approach is. This "
  "note proposes a single addition, an evaluative layer that scores the credibility of each country's "
  "quantum-safe governance against defined criteria, drawn from the comparative-evaluation method that the "
  "assessment literature has converged on and adapted to the Atlas as a standing, auditable base. The aim is "
  "one justified view, not a second instrument: it reuses fields the Atlas already holds, adds a small and "
  "defensible criteria set, and is reported as provisional, consistent with the governance-of-expectations "
  "frame the Atlas rests on.")
for r in ab.runs: r.font.size=Pt(10)

h1("1. What the Atlas measures, and the question it leaves open")
body("The Atlas answers a structural question well. For each country it records the coordination bloc it "
  "migrates with, its standards role, and the governing instrument, its legal status and the obligation it "
  "imposes. That is a map of position: where a country sits in the transition. Position is necessary, and it "
  "is what a first reader needs. It is not sufficient for the question a policy or advisory reader asks next, "
  "which is evaluative rather than descriptive: is this country's approach any good. Two countries can share "
  "a bloc, a role and a binding instrument, and differ sharply in whether their strategy is coherent with "
  "wider policy, whether it is measured, and whether it is governed by a body with the mandate and the skills "
  "to deliver it. The Atlas cannot currently show that difference.")

h1("2. The literature offers a method, and names the Atlas as its base")
body("The QSA methods synthesis re-reads the academic and policy corpus for ways to assess a quantum-safe "
  "transition, and reaches a clear finding: no assessment framework specific to quantum governance yet "
  "exists, a point Paglieri et al. (2025) state directly and build around. Most of the methods the synthesis "
  "assembles are case-level, built to assess an organisation or a venture. One is national in scope and "
  "therefore fits the Atlas's unit of analysis: the comparative policy evaluation that Paglieri et al. (2025) "
  "construct from OECD and World Economic Forum practice. The synthesis notes, in passing, that the Atlas "
  "could provide the auditable base such an evaluation needs, with the added property that it is current. "
  "This proposal takes that observation as its starting point: port the one national method into the Atlas as "
  "a standing layer, so the Atlas becomes the evaluative base the literature says is missing.")

h1("3. The proposed view: a governance-quality layer")
body("The addition is an evaluative reading of each country on six criteria, scored from the documents the "
  "Atlas already cites. Relevance asks whether the national approach is aligned with the threat and carries a "
  "stated transition plan. Coherence asks whether it is integrated with wider cyber and science-and-"
  "technology policy and was built with the state, industry and academia consulted. Effectiveness, judged by "
  "design rather than by outcome, asks whether milestones, measurement and an inventory duty exist. "
  "Efficiency asks whether the instrument mix is proportionate to the stated objective. Governance quality "
  "asks whether a mandate, a lead authority, a steering function and the skills to run the migration are in "
  "place. Impact, again by design, asks whether the approach is structured to reduce the specific exposure, "
  "under a window of risk rather than a single date. Table 1 sets out each criterion, its theoretical "
  "grounding, and the Atlas field it reads from.")

# table 1
rows=[
 ("Relevance","Alignment with the threat; a stated transition plan to quantum-safe cryptography","Comparative evaluation, after OECD and NIST (Paglieri et al., 2025)","Main regulation; migration timeline"),
 ("Coherence","Integration with wider cyber and S&T policy; breadth of consultation across state, industry and academia","OECD and WEF practice (Paglieri et al., 2025); the triple helix (Purohit et al., 2023)","Gov actors; standards participation"),
 ("Effectiveness (by design)","Presence of milestones, measurement and a cryptographic-inventory duty, not yet outcomes","Process evaluation (Vance, 2025); CEPS recommendation on inventories (Pupillo et al., 2025)","Legal status; migration timeline; inventory mandate"),
 ("Efficiency","Proportionality of the instrument mix to the stated objective","Comparative evaluation, after OECD (Paglieri et al., 2025)","Legal status; obligation"),
 ("Governance quality","A mandate, a lead authority, a steering function and migration skills","WEF practice (Paglieri et al., 2025); CEPS capability recommendations (Pupillo et al., 2025)","Gov actors; capacity signals"),
 ("Impact (structured-for)","Whether the approach is structured to reduce the specific exposure, under a Q-period not a Q-day","Impact evaluation (Vance, 2025); the move to a Q-period (Pupillo et al., 2023; 2025)","Target completion; standards role"),
]
t=doc.add_table(rows=1, cols=4); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
lay=OxmlElement('w:tblLayout'); lay.set(qn('w:type'),'fixed'); t._tbl.tblPr.append(lay)
def cell(c,txt,w,bold=False,fill=None,size=9.0):
    c.width=Cm(w); pr=c.paragraphs[0]; pr.paragraph_format.space_after=Pt(1); pr.paragraph_format.line_spacing=1.0
    rr=pr.add_run(txt); rr.bold=bold; rr.font.name="Arial"; rr.font.size=Pt(size); rr.font.color.rgb=INK
    if fill:
        shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),fill); c._tc.get_or_add_tcPr().append(shd)
widths=[3.0,5.6,4.5,3.2]
hdr=t.rows[0].cells
for i,hh in enumerate(["Criterion","What it scores for a country","Theoretical grounding","Reads from"]): cell(hdr[i],hh,widths[i],bold=True,fill="EFECE6")
for row in rows:
    cs=t.add_row().cells
    for i,v in enumerate(row): cell(cs[i],v,widths[i])
cap=doc.add_paragraph(); cr=cap.add_run("Table 1. The evaluative layer: six criteria, their grounding, and the Atlas fields they read from.")
cr.italic=True; cr.font.size=Pt(9); cr.font.color.rgb=MUTED; cap.paragraph_format.space_after=Pt(10)

h1("4. Why this is the right view, and what keeps it honest")
body("The criteria are not invented for the Atlas; they are the OECD and WEF evaluation criteria that "
  "Paglieri et al. (2025) adapted for national quantum strategies, which is why they transfer to a country "
  "instrument cleanly. Three further models sharpen the reading. Vance (2025) separates evaluation into "
  "process, impact and cost-benefit, which is why the layer judges effectiveness and impact by design rather "
  "than by outcomes an emerging field cannot yet show. Marchant et al. (2024) read emerging-technology "
  "governance as a movement from soft law to hard law, which gives the layer a way to read a country's "
  "trajectory, whether its instruments are hardening, rather than a fixed snapshot. The CEPS reports "
  "represent the timing variable as a window of probability, a Q-period rather than a Q-day (Pupillo et al., "
  "2023; 2025), so the impact criterion is scored against an honest distribution rather than a single date.")
body("The whole layer sits inside the frame the Atlas already uses. Governance of expectations treats this "
  "transition as provisional, because the threat's timing is genuinely unknowable (Budde and Konrad, 2019). "
  "An evaluative layer built on that frame does not score a country as finished or correct; it scores how "
  "credible its approach is under irreducible uncertainty, and it expects the score to move. Two disciplines "
  "keep it defensible. Scoring is reconciled by a second reviewer before it is shown, the validation step "
  "Paglieri et al. (2025) build into their method. And the Atlas's existing confidence grade is carried onto "
  "the evaluation, so a country assessed on thin evidence reads as tentative rather than judged.")

h1("5. Why it does not overload the Atlas")
body("The layer is one view, not a second instrument. It reads mostly from fields the Atlas already records, "
  "the regulation, the legal status, the timeline, the actors and the standards role, and needs only the "
  "inventory-duty and capacity signals as genuinely new inputs. It is reported as a quiet, secondary profile "
  "section, not a new map colour, so the structural map remains the hero and colour keeps its single meaning. "
  "It is shown as six transparent sub-scores with their sources, not a composite league-table number, which "
  "is what separates a defensible evaluation from a ranking the evidence cannot support. The limits travel "
  "with it openly: the criteria are qualitative judgements, controlled but not removed by reconciliation, and "
  "they degrade where a national strategy is too nascent to evaluate, the normal condition in this field "
  "(Paglieri et al., 2025). Reported this way, the layer adds the evaluative reading the Atlas lacks without "
  "diluting the structural reading it does well.")

h1("References")
refs=[
 "Budde, B., & Konrad, K. (2019). Tentative governing of fuel cell innovation in a dynamic network of expectations. Research Policy, 48(5), 1098-1112.",
 "Marchant, G., Bazzi, R., Bowman, D., Connor, J., Davis, R. A. III, Kang, E., Konkoly-Thege, K., Liu, D., Lloyd-Jones, S., Manwaring, K., Bennett Moses, L., Wagner, M., & Wastek, S. (2024). Learning from emerging technology governance for guiding quantum technology. UNSW Law & Justice Research Series.",
 "Paglieri, L., Bonomi Savignon, A., Scalabrini, F., & Costumato, L. (2025). Navigating the quantum frontier: Examining government strategy to the next technological revolution. Transforming Government: People, Process and Policy, 19(4).",
 "Purohit, A., Kaur, M., & Venegas-Gomez, A. (2023). Building a quantum-ready ecosystem. IET Quantum Communication.",
 "Pupillo, L., Ferreira, A., Lipiainen, V., & Polito, C. (2023). Quantum technologies and cybersecurity: Technology, governance and policy challenges. Centre for European Policy Studies.",
 "Pupillo, L., et al. (2025). Strengthening the EU transition to a quantum-safe world: Technology, market and governance (Task Force Report). Centre for European Policy Studies.",
 "Vance, A. S. (2025). Cybersecurity and quantum computing: A quantitative analysis proposing a framework for assessing quantum cybersecurity maturity [Doctoral dissertation].",
]
for ref in refs:
    p=doc.add_paragraph(); r=p.add_run(ref); r.font.size=Pt(9.5)
    p.paragraph_format.left_indent=Cm(1.0); p.paragraph_format.first_line_indent=Cm(-1.0); p.paragraph_format.space_after=Pt(5)

doc.save(OUT)
print("saved:", OUT, "| paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables))
