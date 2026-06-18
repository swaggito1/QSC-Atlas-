#!/usr/bin/env python3
# Theoretical-claim verification log for the QSC Atlas companion report.
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/sessions/gracious-kind-wozniak/mnt/qsc-atlas/docs/QSC_Atlas_Theoretical_Verification_Log.docx"
INK = RGBColor(0x1A, 0x1A, 0x1A); MUTED = RGBColor(0x5C, 0x5C, 0x5C)

doc = Document()
n = doc.styles["Normal"]; n.font.name = "Georgia"; n.font.size = Pt(10); n.font.color.rgb = INK
h = doc.styles["Heading 1"]; h.font.name = "Arial"; h.font.size = Pt(14); h.font.color.rgb = INK

sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE; sec.page_width = Cm(29.7); sec.page_height = Cm(21.0)
for m in ("top_margin","bottom_margin","left_margin","right_margin"): setattr(sec, m, Cm(1.6))

doc.add_heading("Theoretical-claim verification log", level=1)
p = doc.add_paragraph(); r = p.add_run(
  "Companion to the QSC Atlas report. Each theoretical attribution was checked against what the cited "
  "source actually argues, on the same basis as the earlier publication fact-check. Verdicts: Supported "
  "(the source backs the claim as used), Tightened (supported, wording made precise), Misattributed (the "
  "source does not support the claim; corrected), Partially supported (reframed so the citation carries "
  "only what it can). All corrections are applied in the current report.")
r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = MUTED

cols = [("Source", 3.4), ("Claim as used in the report", 6.2), ("What the source actually supports", 7.6),
        ("Verdict", 2.7), ("Action taken", 6.6)]
rows = [
 ("Buthe & Mattli (2011)",
  "First-mover advantage, institutional access and capacity determine whose preferences become global rules.",
  "The New Global Rulers: early participation lets actors shape standards; influence rests on timely information and a single voice (institutional complementarity), not raw economic power.",
  "Supported",
  "Kept, and verified against the source text ('who wins and who loses when standard-setting takes place'; institutional complementarity through timely information and a single voice). Now also grounds the standard-maker / contextualiser / taker vocabulary (see Blind row)."),
 ("Yates & Murphy (2019)",
  "A century of rule-making in which de facto global norms emerge without formal imposition.",
  "Engineering Rules: the first global history of voluntary consensus standardisation since 1880; the actors it calls standardisers make standards that become de facto global norms.",
  "Tightened, verified",
  "Reworded to 'voluntary consensus standardisation'; 'market incumbency' removed. Verified against the source text."),
 ("Blind (2013)",
  "Original draft: a 'typology of standardisation strategies' supplying the standard-maker / contextualiser / taker vocabulary.",
  "Blind gives a fourfold typology of standards FUNCTIONS (compatibility, quality, variety reduction, information), not a typology of maker/taker strategies.",
  "Misattributed",
  "Corrected. Blind now cited for the functional typology; the standard-maker / contextualiser / taker vocabulary is reattributed to the power asymmetry in Buthe & Mattli (2011). NB: the same misattribution appears in the thesis (Ch 6.1.3) and should be fixed there too."),
 ("Bradford (2020)",
  "The Brussels Effect, with the direction reversing in cryptography.",
  "The Brussels Effect (EU regulation diffusing globally through market access) is Bradford's. The reversal in cryptography is the report's own adaptation, not Bradford's claim.",
  "Supported",
  "Kept. The reversal stays attributed to the report, not to Bradford."),
 ("Rodrik (2018)",
  "Original draft: trade and market-access consequences that follow from divergent technical regulation.",
  "Trade agreements have moved beyond tariffs into regulatory rules with distributional, rent-seeking consequences. Rodrik does not address technical standards or market access directly.",
  "Partially supported",
  "Reframed (the comment flag); after Swann was removed, the market-access point is stated analytically and Rodrik carries only the trade-agreements-into-regulation point."),
 ("Swann (2010)",
  "Used to carry the standards-as-market-access link.",
  "The OECD paper is on topic, but the researcher queried the citation as inaccurate for this use and asked for its removal.",
  "Removed",
  "Deleted from the report at the researcher's instruction. The market-access point is retained as an analytical statement without a citation."),
 ("Borup et al. (2006); Konrad (2006); Budde & Konrad (2019)",
  "Governance of expectations: expectations coordinate action under uncertainty and harden into commitment while remaining revisable.",
  "Foundational sociology and governance of expectations; Budde & Konrad on tentative governing in a dynamic network of expectations.",
  "Supported",
  "Kept."),
 ("Star (1999)",
  "Cryptographic protocols are infrastructure, visible only on breakdown.",
  "'The ethnography of infrastructure': infrastructure is invisible until breakdown and embedded in an installed base.",
  "Supported",
  "Kept."),
 ("Carlsson (2006)",
  "Internationalisation of innovation systems frames the cross-border circulation.",
  "A survey of the internationalisation of innovation systems.",
  "Supported",
  "Kept."),
 ("Institutional sources (NIST 2024; NIS2; DORA; Rec (EU) 2024/1101; NIS CG 2025)",
  "FIPS 203/204/205 finalised in 2024; EU binding and soft instruments; the coordinated roadmap and its milestones.",
  "Confirmed against the official documents during the Atlas build (NIST, the EU instruments, the NIS Cooperation Group roadmap).",
  "Supported",
  "Kept."),
 ("Terminology: standard-setter to standard-maker",
  "The role label 'standard-setter' is changed to 'standard-maker' at the researcher's request.",
  "Neither source uses 'maker': Yates & Murphy (2019) use 'standardisers' and 'standard setting', and Buthe & Mattli (2011) use 'standard-setting'. The term is the Atlas's own, grounded in their concepts (standards made through consensus; first-mover advantage deciding whose making prevails).",
  "Applied",
  "Labels and prose updated across the report, the model document, the figures and the live-site role badge; internal data keys left as 'setter' for stability."),
]

table = doc.add_table(rows=1, cols=len(cols)); table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'),'fixed'); table._tbl.tblPr.append(lay)
def cell(c, text, w, bold=False, size=8.5, fill=None):
    c.width = Cm(w); pr=c.paragraphs[0]; pr.paragraph_format.space_after=Pt(1); pr.paragraph_format.line_spacing=1.0
    rr=pr.add_run(text); rr.bold=bold; rr.font.name="Arial"; rr.font.size=Pt(size); rr.font.color.rgb=INK
    if fill:
        shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),fill); c._tc.get_or_add_tcPr().append(shd)
hdr=table.rows[0].cells
for i,(name,w) in enumerate(cols): cell(hdr[i],name,w,bold=True,size=9,fill="EFECE6")
for row in rows:
    cs=table.add_row().cells
    for i,(name,w) in enumerate(cols): cell(cs[i],row[i],w)

doc.save(OUT)
print("saved:", OUT, "| rows:", len(rows))
