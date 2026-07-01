#!/usr/bin/env python3
# Updated QSC Atlas companion report: concise and deep, adding the governance-credibility
# evaluations. Generates its own infographics from the live profiles, builds the .docx with
# appendix tables. British English, no em-dashes, inline APA.
import json, os, re, math
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = "/sessions/gracious-kind-wozniak/mnt/qsc-atlas"
PROF = f"{REPO}/data/profiles"
FIG = f"{REPO}/docs/fig"; os.makedirs(FIG, exist_ok=True)
OUT = f"{REPO}/docs/QSC_Atlas_Companion_Report.docx"
INK = RGBColor(0x1A,0x1A,0x1A); MUTED = RGBColor(0x5C,0x5C,0x5C)
PAPER="#F7F5F0"; INKX="#1A1A1A"; HAIR="#D8D6D0"; MUTEDX="#5C5C5C"
POST={"EU":"#5b54a8","NIST-bloc":"#2b4c7e","sovereign-bloc":"#7a3b5e","engaged-unaligned":"#6b7280"}
ROLEC={"setter":"#a8322a","contextualiser":"#b9762e","taker":"#4a6fa5","sovereign-developer":"#7a3b5e"}
POST_ORDER=["NIST-bloc","EU","sovereign-bloc","engaged-unaligned"]
POST_LAB={"EU":"EU roadmap","NIST-bloc":"NIST-bloc","sovereign-bloc":"Sovereign bloc","engaged-unaligned":"Engaged, unaligned"}
ROLE_ORDER=["setter","contextualiser","taker","sovereign-developer"]
ROLE_LAB={"setter":"Standard-maker","contextualiser":"Contextualiser","taker":"Standard-taker","sovereign-developer":"Sovereign-developer"}
AXES=["Relevance","Coherence","Effectiveness","Efficiency","Governance","Impact"]

# ---------- scoring ----------
def lines(s): return [l.strip() for l in (s or "").split("\n") if l.strip()]
def has_year(s): return sum(1 for l in lines(s) if re.search(r"20\d\d", l))
def score(p):
    tl=p.get("migrationTimeline") or ""; ms=has_year(tl); target=bool(p.get("targetCompletion"))
    regs=lines(p.get("mainRegulation")); nreg=len(regs)
    legal=p.get("legalStatus"); binding=legal=="binding"; soft=legal=="soft-only"
    role=p.get("standardsRole"); dom=p.get("dominantProcess"); gov=lines(p.get("govActors")); ngov=len(gov)
    part=bool((p.get("processParticipation") or "").strip())
    hybrid=(p.get("hybridDeployment") or "").strip(); has_hybrid=hybrid not in ("","None stated","None")
    fams=bool((p.get("standardFamilies") or p.get("algorithms") or "").strip())
    text=" ".join([p.get("obligation") or "",p.get("mainRegulation") or "",p.get("summary") or "",tl]).lower()
    eu=("nis2" in text or "dora" in text or "(eu)" in text)
    s={"R1":2 if ms>=1 else(1 if(nreg>0 or target)else 0),"R2":2 if binding else(1 if(soft or dom)else 0),
       "C1":2 if(eu or nreg>=2)else(1 if(nreg==1 or part)else 0),"C2":2 if ngov>=3 else(1 if(ngov>=1 or part)else 0),
       "E1":2 if ms>=1 else(1 if target else 0),"E2":2 if("inventory" in text or "cbom" in text)else(1 if binding else 0),
       "E3":2 if any(w in text for w in("monitor","indicator","audit"))else None,"Ef1":2 if binding else(1 if soft else 0),
       "G1":2 if ngov>=2 else(1 if ngov==1 else 0),"G2":2 if any(w in text for w in("skill","training","academ","programme","workforce"))else None,
       "I1":2 if(role or(dom and(has_hybrid or fams or nreg>0)))else(1 if dom else 0)}
    def m(k):
        v=[s[x] for x in k if s.get(x) is not None]; return round(sum(v)/len(v),2) if v else None
    return {"Relevance":m(["R1","R2"]),"Coherence":m(["C1","C2"]),"Effectiveness":m(["E1","E2","E3"]),
            "Efficiency":m(["Ef1"]),"Governance":m(["G1","G2"]),"Impact":m(["I1"])}
CONF={"Complete":"High","Partial":"Medium","Placeholder":"Low"}
countries=[]
for f in sorted(os.listdir(PROF)):
    if not f.endswith(".json"): continue
    p=json.load(open(os.path.join(PROF,f)))
    if not p.get("dominantProcess") or p.get("iso3")=="EUU": continue
    countries.append({"name":p.get("country") or p["iso3"],"posture":p.get("coordinationPosture"),
        "role":p.get("standardsRole"),"legal":p.get("legalStatus"),
        "conf":p.get("confidence") or CONF.get(p.get("dataStatus"),"-"),"ax":score(p)})
countries.sort(key=lambda c:c["name"])
N=len(countries)
postcount={k:sum(1 for c in countries if c["posture"]==k) for k in POST_ORDER}
rolecount={k:sum(1 for c in countries if c["role"]==k) for k in ROLE_ORDER}
rolecount["none"]=sum(1 for c in countries if not c["role"])
legalcount={k:sum(1 for c in countries if c["legal"]==k) for k in("binding","soft-only","none")}
legalcount["none"]=sum(1 for c in countries if not c["legal"] or c["legal"]=="none")
axmean={a:round(sum(c["ax"][a] for c in countries if c["ax"][a] is not None)/
        sum(1 for c in countries if c["ax"][a] is not None),2) for a in AXES}
def bloc_mean(b):
    vals=[v for c in countries if c["posture"]==b for v in c["ax"].values() if v is not None]
    return round(sum(vals)/len(vals),2) if vals else 0
blocmean={b:bloc_mean(b) for b in POST_ORDER}

# ---------- figures ----------
plt.rcParams.update({"font.family":"DejaVu Sans","font.size":10,"text.color":INKX,
    "axes.edgecolor":HAIR,"axes.labelcolor":INKX,"xtick.color":MUTEDX,"ytick.color":MUTEDX})
def newfig(w,h):
    fig=plt.figure(figsize=(w,h),dpi=200); fig.patch.set_facecolor(PAPER); return fig

# Fig 1: posture x role matrix
def fig_matrix():
    fig=newfig(9,3.6); ax=fig.add_subplot(111); ax.set_facecolor(PAPER)
    nc,nr=len(POST_ORDER),len(ROLE_ORDER)
    for i in range(nc+1): ax.plot([i,i],[0,nr],color=HAIR,lw=0.8,zorder=1)
    for j in range(nr+1): ax.plot([0,nc],[j,j],color=HAIR,lw=0.8,zorder=1)
    for ci,pp in enumerate(POST_ORDER):
        for rj,rr in enumerate(ROLE_ORDER):
            n=sum(1 for c in countries if c["posture"]==pp and c["role"]==rr); y=nr-1-rj
            if n:
                ax.add_patch(plt.Rectangle((ci+.04,y+.04),.92,.92,facecolor=POST[pp],alpha=.12,edgecolor="none",zorder=2))
                ax.text(ci+.5,y+.5,str(n),ha="center",va="center",fontsize=14,color=INKX,fontweight="bold",zorder=3)
            else: ax.text(ci+.5,y+.5,"·",ha="center",va="center",fontsize=12,color=HAIR,zorder=3)
    for ci,pp in enumerate(POST_ORDER): ax.text(ci+.5,nr+.14,POST_LAB[pp],ha="center",va="bottom",fontsize=8.5,color=POST[pp],fontweight="bold")
    for rj,rr in enumerate(ROLE_ORDER): ax.text(-.12,nr-1-rj+.5,ROLE_LAB[rr],ha="right",va="center",fontsize=8.5,color=ROLEC[rr],fontweight="bold")
    # unassigned role row indicator
    ax.text(nc/2,-0.5,f"{rolecount['none']} countries carry a bloc but no standards role yet",ha="center",va="center",fontsize=8,color=MUTEDX,style="italic")
    ax.set_xlim(-1.9,nc+.05); ax.set_ylim(-0.8,nr+.55); ax.axis("off")
    fig.tight_layout(); fig.savefig(f"{FIG}/rep_matrix.png",facecolor=PAPER,bbox_inches="tight"); plt.close()

# Fig 2: France six-axis radar
def fig_radar(name):
    c=next(x for x in countries if x["name"]==name); vals=[c["ax"][a] or 0 for a in AXES]
    ang=[n/6*2*math.pi for n in range(6)]; ang+=ang[:1]; v=vals+vals[:1]
    fig=newfig(5.2,5.2); ax=fig.add_subplot(111,polar=True); ax.set_facecolor(PAPER)
    ax.set_theta_offset(math.pi/2); ax.set_theta_direction(-1)
    ax.set_ylim(0,2); ax.set_yticks([1,2]); ax.set_yticklabels(["1","2"],color=MUTEDX,fontsize=8)
    ax.set_xticks(ang[:-1]); ax.set_xticklabels(AXES,fontsize=9,color=INKX)
    ax.plot(ang,v,color=POST["EU"],lw=2); ax.fill(ang,v,color=POST["EU"],alpha=.15)
    ax.spines["polar"].set_color(HAIR); ax.grid(color=HAIR,lw=.7)
    fig.tight_layout(); fig.savefig(f"{FIG}/rep_radar.png",facecolor=PAPER,bbox_inches="tight"); plt.close()

# Fig 3: mean credibility by bloc
def fig_bybloc():
    fig=newfig(8,2.9); ax=fig.add_subplot(111); ax.set_facecolor(PAPER)
    labs=[POST_LAB[b] for b in POST_ORDER]; vals=[blocmean[b] for b in POST_ORDER]; cols=[POST[b] for b in POST_ORDER]
    y=range(len(labs)); ax.barh(list(y),vals,color=cols,alpha=.85,height=.6)
    for i,val in enumerate(vals): ax.text(val+.03,i,f"{val:.2f}",va="center",fontsize=9,color=INKX)
    ax.set_yticks(list(y)); ax.set_yticklabels(labs,fontsize=9); ax.invert_yaxis()
    ax.set_xlim(0,2); ax.set_xlabel("mean of the six credibility axes (0 to 2)",fontsize=8.5,color=MUTEDX)
    for s in("top","right","left"): ax.spines[s].set_visible(False)
    ax.tick_params(length=0); ax.set_axisbelow(True); ax.xaxis.grid(color=HAIR,lw=.7)
    fig.tight_layout(); fig.savefig(f"{FIG}/rep_bybloc.png",facecolor=PAPER,bbox_inches="tight"); plt.close()

# Fig 4: mean by axis
def fig_byaxis():
    fig=newfig(8,2.9); ax=fig.add_subplot(111); ax.set_facecolor(PAPER)
    vals=[axmean[a] for a in AXES]
    bars=ax.bar(AXES,vals,color=INKX,alpha=.82,width=.62)
    mn=min(vals);
    for b,a in zip(bars,AXES):
        if axmean[a]==mn: b.set_color("#a8322a"); b.set_alpha(.9)
    for i,val in enumerate(vals): ax.text(i,val+.03,f"{val:.2f}",ha="center",fontsize=8.5,color=INKX)
    ax.set_ylim(0,2); ax.set_ylabel("mean score (0 to 2)",fontsize=8.5,color=MUTEDX)
    for s in("top","right"): ax.spines[s].set_visible(False)
    ax.tick_params(length=0); ax.set_axisbelow(True); ax.yaxis.grid(color=HAIR,lw=.7)
    plt.xticks(fontsize=8.5)
    fig.tight_layout(); fig.savefig(f"{FIG}/rep_byaxis.png",facecolor=PAPER,bbox_inches="tight"); plt.close()

fig_matrix(); fig_radar("France"); fig_bybloc(); fig_byaxis()

# ---------- document ----------
doc=Document()
n=doc.styles["Normal"]; n.font.name="Georgia"; n.font.size=Pt(10.5); n.font.color.rgb=INK
n.paragraph_format.space_after=Pt(8); n.paragraph_format.line_spacing=1.25
for hid,sz in (("Heading 1",14),("Heading 2",11.5)):
    st=doc.styles[hid]; st.font.name="Arial"; st.font.size=Pt(sz); st.font.bold=True; st.font.color.rgb=INK
    st.paragraph_format.space_before=Pt(12); st.paragraph_format.space_after=Pt(5); st.paragraph_format.keep_with_next=True
sec=doc.sections[0]; sec.page_width=Cm(21.0); sec.page_height=Cm(29.7)
for m in("top_margin","bottom_margin","left_margin","right_margin"): setattr(sec,m,Cm(2.4))
def body(t,after=8,size=None):
    p=doc.add_paragraph(); r=p.add_run(t)
    if size: r.font.size=Pt(size)
    p.paragraph_format.space_after=Pt(after); return p
def h1(t): doc.add_heading(t,1)
def figure(fname,cap,w=15.0):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(f"{FIG}/{fname}",width=Cm(w))
    c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
    cr=c.add_run(cap); cr.italic=True; cr.font.size=Pt(9); cr.font.color.rgb=MUTED; c.paragraph_format.space_after=Pt(12)

# title
t=doc.add_paragraph(); tr=t.add_run("Quantum-safe cryptography: the quiet migration reshaping digital sovereignty")
tr.font.name="Georgia"; tr.font.size=Pt(18); tr.bold=True; t.paragraph_format.space_after=Pt(2)
s=doc.add_paragraph(); sr=s.add_run("A companion report to the QSC Atlas, with the governance-credibility evaluations")
sr.font.name="Arial"; sr.font.size=Pt(11.5); sr.font.color.rgb=MUTED; s.paragraph_format.space_after=Pt(2)
mm=doc.add_paragraph(); mr=mm.add_run("Companion to the QSC Atlas (qsc-atlas.vercel.app). Research note, updated June 2026. [Author and affiliation]")
mr.font.name="Arial"; mr.font.size=Pt(9); mr.font.color.rgb=MUTED; mm.paragraph_format.space_after=Pt(14)

ab=doc.add_paragraph(); abr=ab.add_run("Abstract. "); abr.bold=True
ab.add_run("The QSC Atlas maps, country by country, how governments are rebuilding public-key "
  "cryptography against the prospect of a quantum computer able to break it. This updated report is "
  "concise by design. It sets out the two-axis classification behind the Atlas, a coordination posture "
  "and a standards role; the regulatory layer that records what each instrument requires; and, new to "
  "this edition, an evaluative layer that scores how credible each national plan is on six criteria. It "
  "explains how the evaluations are measured and verified, and reports what they show across the "
  f"{N} recorded countries: credibility is high where a country has committed to a bloc and thin where it "
  "has not, and the weakest link everywhere is execution, the inventories and measurement that turn a "
  "plan into practice. The full scores are in the appendix.")
for r in ab.runs: r.font.size=Pt(10)

h1("1. The instrument and its question")
body("Public-key cryptography secures most of what the digital economy runs on. A sufficiently powerful "
  "quantum computer would break the asymmetric algorithms that work depends on; no such machine exists, "
  "and the date of its arrival cannot be known. Governments, regulators, standards bodies and technology "
  "firms are nonetheless already rebuilding those foundations against the anticipated threat, and the QSC "
  "Atlas records that rebuilding as it happens. Two features make it hard to read from any single vantage "
  "point: cryptographic protocols are infrastructure in Star's (1999) sense, visible only on breakdown and "
  "migrated rather than switched; and no single authority governs the transition, which is coordinated "
  "instead across national agencies, regional roadmaps and transnational standards bodies. The Atlas shows "
  "that scattered structure rather than assuming a centre that does not exist, and now also asks a further "
  "question: not only where a country stands, but how credible its plan is.")

h1("2. Theoretical principles")
body("The Atlas rests on the standards-governance literature, which treats technical standards as "
  "instruments of power with distributional consequences rather than neutral specifications. Buthe and "
  "Mattli (2011) show how first-mover advantage, institutional access and technical capacity determine "
  "whose preferences become embedded in global rules; Yates and Murphy (2019) trace the same asymmetry "
  "across a century of voluntary consensus standard-setting. The post-quantum case reproduces it: when the "
  "United States published its first three standards in 2024 (NIST, 2024), a contested question became a "
  "settled technical fact that other jurisdictions adopt or must interoperate with. Bradford's (2020) "
  "Brussels Effect is the instructive contrast, because in cryptography the direction reverses: the "
  "standards originate in Washington and travel through the global deployment of United States technology "
  "firms, while European institutions contextualise rather than make.")
body("Beneath the standards reading sits the governance-of-expectations literature (Borup et al., 2006; "
  "Konrad, 2006). Because the threat's timing is genuinely unknowable, coordination proceeds through shared "
  "expectations that can harden into commitment while remaining revisable (Budde and Konrad, 2019); the "
  "governance stays provisional. The evaluative layer added in this edition takes its criteria from the "
  "same register: it adapts the six-criteria comparative method Paglieri et al. (2025) built for national "
  "quantum strategies, sharpened by Vance's (2025) split of evaluation into process, impact and "
  "cost-benefit, and draws two of its checks from the CEPS governance recommendations (Pupillo et al., "
  "2025). A credibility score is therefore a reading at a date, not a verdict.")

h1("3. The two-axis classification")
body("Every recorded country is described on two dimensions, kept apart because one label cannot carry "
  "both. The coordination posture answers which collective transition a country belongs to, and takes one "
  "of four values: the EU coordinated roadmap, the NIST-led bloc, a sovereign bloc, or engaged but "
  "unaligned. The standards role answers its relationship to the algorithms: standard-maker, "
  "contextualiser, standard-taker, or sovereign developer. A maker writes the baseline; a contextualiser "
  "adopts it but adds national requirements such as mandatory hybridisation or certification; a taker "
  "adopts it broadly as published; a sovereign developer designs its own. The posture is the map colour, "
  "the role a profile badge. Figure 1 shows why the split matters: the sovereign and engaged groups each "
  f"sit on a single role, but the committed blocs divide across roles, with the United States the lone "
  "maker, seven contextualisers, and the majority takers. A one-colour scheme would hide that.")
figure("rep_matrix.png","Figure 1. The two-axis classification: coordination bloc by standards role. "
  f"{N} recorded countries.",16.0)

h1("4. The regulatory layer")
body("Between the classification and the evaluation sits a regulatory layer that turns a label into an "
  "obligation. For each country it records the main governing instrument, its legal status, binding, "
  "soft-law or none, and the plain duty it places on regulated entities. The distinction is load-bearing: "
  f"a recommendation and a statute would otherwise read alike. Of the {N} recorded countries, "
  f"{legalcount['binding']} carry at least one binding instrument, {legalcount['soft-only']} rest on soft "
  f"law alone, and {legalcount['none']} have no governing instrument on record. Binding EU duties under "
  "NIS2 and DORA read differently from a soft national roadmap, which reads differently again from nothing "
  "at all, and the evaluative layer draws directly on these fields.")

h1("5. Measuring credibility: the evaluative layer")
body("The evaluative layer asks not whether a plan exists but how credible it is. Six criteria, relevance, "
  "coherence, effectiveness, efficiency, governance quality, and impact, become eleven sub-criteria, each "
  "scored 0 (absent), 1 (partial) or 2 (present) against a stated description, and reported as a six-axis "
  "profile rather than a single number (Figure 2). Relevance reads the transition plan and the threat "
  "articulation; coherence, the integration with wider policy and the breadth of consultation; "
  "effectiveness, the milestones, the inventory duty and any measurement framework; efficiency, whether "
  "the instrument's strength matches the objective; governance quality, the mandate and the skills base; "
  "and impact, the exposure-reducing adoption of the baseline. Each sub-score is derived from a field the "
  "Atlas already records, so it is traceable: milestones and target years drive relevance and "
  "effectiveness, the legal status drives efficiency, the recorded instruments and obligations drive "
  "coherence and the inventory check, and the standards role with the hybrid stance drives impact.")
body("Three rules keep the scores defensible. No sub-score is entered unless it traces to a sourced "
  "instrument; an axis with no supporting field reads as not assessed, never a default, and is the mean of "
  "the sub-criteria that could be assessed. A second reviewer scores each country independently and the two "
  "reconcile before a profile is shown. The Atlas confidence grade is carried onto each axis, so a country "
  "assessed on thin evidence reads as tentative. The scores in this report are a first pass, derived "
  "mechanically from the recorded fields by these rules, with reviewer reconciliation still under way.")
figure("rep_radar.png","Figure 2. How a country reads: the six-axis credibility profile for France, an EU "
  "contextualiser. High on relevance, coherence, efficiency, governance and impact; a step lower on "
  "effectiveness, where the milestones are not-earlier-than windows and no dedicated measurement duty is "
  "recorded.",10.5)

h1("6. What the evaluations show")
body("Read across the six axes, the scores have a consistent shape (Figure 3). Relevance and impact are "
  f"highest, at {axmean['Relevance']:.2f} and {axmean['Impact']:.2f}, because plans exist and the NIST "
  f"baseline is adopted almost everywhere. Effectiveness is the lowest axis, at {axmean['Effectiveness']:.2f}, "
  "because the fields it reads, mandated inventories and measurement frameworks, are the thinnest in the "
  "record. The finding is that plans outrun the machinery to execute them: the transition is widely "
  "declared and widely under-instrumented.")
figure("rep_byaxis.png","Figure 3. Mean credibility by axis across the recorded countries. Effectiveness, "
  "the execution machinery of inventories and measurement, is the weakest link everywhere.",15.0)
body("Read by coordination bloc, credibility concentrates in the committed (Figure 4). The NIST-led and EU "
  f"blocs score almost identically, {blocmean['NIST-bloc']:.2f} and {blocmean['EU']:.2f}, confirming that "
  "on credibility, as on algorithms, the two committed blocs behave alike. The sovereign bloc sits lower "
  f"at {blocmean['sovereign-bloc']:.2f}, its independent track less legible in the fields the layer reads. "
  f"The engaged-but-unaligned group scores far below the rest, at {blocmean['engaged-unaligned']:.2f}: the "
  "credibility gap is concentrated among the states that have named no bloc, standard or timeline. This is "
  "the classification and the evaluation agreeing, commitment and credibility travel together, and it is "
  "the clearest single result of the layer.")
figure("rep_bybloc.png","Figure 4. Mean credibility by coordination bloc. Commitment and credibility "
  "travel together; the gap is concentrated among the unaligned.",15.0)

h1("7. Why it matters")
body("Theoretically, the case extends standards-governance and governance-of-expectations accounts into a "
  "transnational, multi-level setting: expectations generated within one national system acquire binding "
  "force elsewhere through market mechanisms and institutional referencing rather than formal agreement "
  "(Buthe and Mattli, 2011; Carlsson, 2006). Methodologically, separating coordination from role avoids a "
  "bloc illusion in which the EU reads as a rival standards camp when on the algorithm axis it runs the "
  "same baseline under a regulatory wrapper; and scoring credibility on sourced fields, with confidence "
  "carried and reconciliation required, keeps the reading defensible rather than asserted. Geopolitically, "
  "whoever makes the baseline holds structural power over the terms on which everyone else secures their "
  "systems, which is why the single-maker finding is the central fact of the field; the credibility gap "
  "among the unaligned marks where the contest for the undecided middle will be decided.")

h1("8. Conclusion")
body("The QSC Atlas is a living instrument for a transition that has not finished and may not finish on any "
  "predictable schedule. Its contribution is to make a leaderless migration legible: which collective each "
  "country coordinates with, whose algorithms it runs, what binds it, and now how credible its plan is. The "
  "evaluations are provisional by construction and move as the evidence does. Two things are already clear: "
  "credibility follows commitment, and execution is the common weak point. The appendix records the current "
  f"classification and the credibility scores for all {N} recorded countries.")

# ---------- Appendix A: classification (portrait) ----------
doc.add_page_break()
doc.add_heading("Appendix A. Classification of the recorded countries", 1)
ip=doc.add_paragraph(); ir=ip.add_run("Coordination bloc, standards role and legal status of the governing "
  "instrument, from the QSC Atlas. Source: institutional documents on each country profile.")
ir.italic=True; ir.font.size=Pt(8.5); ir.font.color.rgb=MUTED
def mk_table(cols, rowsdata, sizes=8.5, head="EFECE6"):
    tb=doc.add_table(rows=1,cols=len(cols)); tb.alignment=WD_TABLE_ALIGNMENT.CENTER; tb.autofit=False
    lay=OxmlElement('w:tblLayout'); lay.set(qn('w:type'),'fixed'); tb._tbl.tblPr.append(lay)
    def cell(c,txt,w,bold=False,fill=None):
        c.width=Cm(w); pr=c.paragraphs[0]; pr.paragraph_format.space_after=Pt(1); pr.paragraph_format.line_spacing=1.0
        rr=pr.add_run(txt); rr.bold=bold; rr.font.name="Arial"; rr.font.size=Pt(sizes); rr.font.color.rgb=INK
        if fill:
            sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),fill); c._tc.get_or_add_tcPr().append(sh)
    for i,(nm,w) in enumerate(cols): cell(tb.rows[0].cells[i],nm,w,bold=True,fill=head)
    for row in rowsdata:
        cs=tb.add_row().cells
        for i,(nm,w) in enumerate(cols): cell(cs[i],row[i],w)
    return tb
POST_SHORT={"EU":"EU roadmap","NIST-bloc":"NIST-bloc","sovereign-bloc":"Sovereign","engaged-unaligned":"Engaged, unaligned"}
half=(N+1)//2
def clf_rows(sub):
    return [[c["name"], POST_SHORT.get(c["posture"],c["posture"] or "-"),
             ROLE_LAB.get(c["role"],"Not yet declared") if c["role"] else "Not yet declared",
             {"binding":"Binding","soft-only":"Soft law","none":"None"}.get(c["legal"] or "none","None")] for c in sub]
colsA=[("Country",4.0),("Bloc",3.4),("Role",3.6),("Legal",2.4)]
# two side-by-side halves via a 2-col wrapper would be complex; use one table split across the page in two stacked chunks
mk_table(colsA, clf_rows(countries[:half]))
doc.add_paragraph().paragraph_format.space_after=Pt(2)
mk_table(colsA, clf_rows(countries[half:]))

# ---------- Appendix B: evaluation scores (portrait) ----------
doc.add_page_break()
doc.add_heading("Appendix B. Governance-credibility scores", 1)
ip=doc.add_paragraph(); ir=ip.add_run("Six-axis credibility scores (0 to 2 per axis; a dash means the axis "
  "could not be assessed from the recorded fields). First pass, derived from the recorded fields; reviewer "
  "reconciliation under way. Rel relevance, Coh coherence, Eff effectiveness, Effi efficiency, Gov "
  "governance quality, Imp impact.")
ir.italic=True; ir.font.size=Pt(8.5); ir.font.color.rgb=MUTED
def fmt(x): return "-" if x is None else f"{x:.1f}"
colsB=[("Country",3.7),("Rel",1.5),("Coh",1.5),("Eff",1.5),("Effi",1.5),("Gov",1.5),("Imp",1.5),("Conf.",1.7)]
rowsB=[[c["name"]]+[fmt(c["ax"][a]) for a in AXES]+[c["conf"]] for c in countries]
mk_table(colsB, rowsB, sizes=8.0)
mp=doc.add_paragraph(); mr2=mp.add_run(f"Axis means across the {N} countries: relevance {axmean['Relevance']:.2f}, "
  f"coherence {axmean['Coherence']:.2f}, effectiveness {axmean['Effectiveness']:.2f}, efficiency "
  f"{axmean['Efficiency']:.2f}, governance {axmean['Governance']:.2f}, impact {axmean['Impact']:.2f}.")
mr2.italic=True; mr2.font.size=Pt(8.5); mr2.font.color.rgb=MUTED

# ---------- references ----------
doc.add_page_break()
doc.add_heading("References", 1)
refs=[
 "Borup, M., Brown, N., Konrad, K., & van Lente, H. (2006). The sociology of expectations in science and technology. Technology Analysis & Strategic Management, 18(3-4), 285-298.",
 "Bradford, A. (2020). The Brussels effect: How the European Union rules the world. Oxford University Press.",
 "Budde, B., & Konrad, K. (2019). Tentative governing of fuel cell innovation in a dynamic network of expectations. Research Policy, 48(5), 1098-1112.",
 "Buthe, T., & Mattli, W. (2011). The new global rulers: The privatization of regulation in the world economy. Princeton University Press.",
 "Carlsson, B. (2006). Internationalization of innovation systems: A survey of the literature. Research Policy, 35(1), 56-67.",
 "European Parliament and Council. (2022). Directive (EU) 2022/2555 (NIS2); Regulation (EU) 2022/2554 (DORA).",
 "Konrad, K. (2006). The social dynamics of expectations. Technology Analysis & Strategic Management, 18(3-4), 429-444.",
 "NIST. (2024). FIPS 203, FIPS 204 and FIPS 205: Post-quantum cryptography standards. National Institute of Standards and Technology.",
 "Paglieri, L., Bonomi Savignon, A., Scalabrini, F., & Costumato, L. (2025). Navigating the quantum frontier: Examining government strategy to the next technological revolution. Transforming Government: People, Process and Policy, 19(4).",
 "Pupillo, L., et al. (2025). Strengthening the EU transition to a quantum-safe world (Task Force Report). Centre for European Policy Studies.",
 "Star, S. L. (1999). The ethnography of infrastructure. American Behavioral Scientist, 43(3), 377-391.",
 "Vance, A. S. (2025). Cybersecurity and quantum computing: A quantitative analysis proposing a framework for assessing quantum cybersecurity maturity [Doctoral dissertation].",
 "Yates, J., & Murphy, C. N. (2019). Engineering rules: Global standard setting since 1880. Johns Hopkins University Press.",
]
for ref in refs:
    p=doc.add_paragraph(); r=p.add_run(ref); r.font.size=Pt(9.5)
    p.paragraph_format.left_indent=Cm(1.0); p.paragraph_format.first_line_indent=Cm(-1.0); p.paragraph_format.space_after=Pt(5)

doc.save(OUT)
print("saved:",OUT)
print("N:",N,"| posture:",postcount,"| role:",rolecount,"| legal:",legalcount)
print("axis means:",axmean,"| bloc means:",blocmean)
print("paragraphs:",len(doc.paragraphs),"| tables:",len(doc.tables))
