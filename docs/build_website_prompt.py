#!/usr/bin/env python3
# The QSC Atlas website rebuild: a precise, scroll-driven build brief for Claude Code,
# grounded in the ui-ux-pro-max design skill and merged with the evaluative-layer
# methodology. British English, no em-dashes, monochrome discipline. python-docx.
import json, os, re
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import font_manager
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = "/sessions/gracious-kind-wozniak/mnt/qsc-atlas"
PROF = f"{REPO}/data/profiles"
FIGD = f"{REPO}/docs/fig"; os.makedirs(FIGD, exist_ok=True)
OUT = f"{REPO}/docs/QSC_Atlas_Website_Rebuild_Brief.docx"
INK = RGBColor(0x1A,0x1A,0x1A); MUTED = RGBColor(0x5C,0x5C,0x5C); WASH="EFECE6"; CODEBG="F1EFE9"
PAPER="#F7F5F0"; INKHEX="#1A1A1A"; HAIR="#D8D6D0"
POSTURE={"EU":"#5b54a8","NIST-bloc":"#2b4c7e","sovereign-bloc":"#7a3b5e","engaged-unaligned":"#6b7280"}
ROLE={"setter":"#a8322a","contextualiser":"#b9762e","taker":"#4a6fa5","sovereign-developer":"#7a3b5e"}

# ---------- load profiles ----------
profs=[]
for f in sorted(os.listdir(PROF)):
    if not f.endswith(".json"): continue
    p=json.load(open(os.path.join(PROF,f)))
    if p.get("dominantProcess") and p.get("iso3")!="EUU": profs.append(p)

# ---------- infographic: posture x role grid ----------
def make_infographic():
    postures=["EU","NIST-bloc","sovereign-bloc","engaged-unaligned"]
    roles=["setter","contextualiser","taker","sovereign-developer"]
    rolelab={"setter":"Standard-maker","contextualiser":"Contextualiser","taker":"Standard-taker","sovereign-developer":"Sovereign-developer"}
    grid={(pp,rr):0 for pp in postures for rr in roles}
    for p in profs:
        pp=p.get("coordinationPosture"); rr=p.get("standardsRole")
        if pp in postures and rr in roles: grid[(pp,rr)]+=1
    fig,ax=plt.subplots(figsize=(9.2,4.3),dpi=200); fig.patch.set_facecolor(PAPER); ax.set_facecolor(PAPER)
    nc,nr=len(postures),len(roles)
    for i in range(nc+1): ax.plot([i,i],[0,nr],color=HAIR,lw=0.8,zorder=1)
    for j in range(nr+1): ax.plot([0,nc],[j,j],color=HAIR,lw=0.8,zorder=1)
    for ci,pp in enumerate(postures):
        for rj,rr in enumerate(roles):
            n=grid[(pp,rr)]; y=nr-1-rj
            if n>0:
                ax.add_patch(plt.Rectangle((ci+0.04,y+0.04),0.92,0.92,facecolor=POSTURE[pp],alpha=0.10,edgecolor="none",zorder=2))
                ax.text(ci+0.5,y+0.5,str(n),ha="center",va="center",fontsize=15,color=INKHEX,zorder=3,fontweight="bold")
            else:
                ax.text(ci+0.5,y+0.5,"·",ha="center",va="center",fontsize=14,color=HAIR,zorder=3)
    for ci,pp in enumerate(postures):
        ax.text(ci+0.5,nr+0.18,pp,ha="center",va="bottom",fontsize=8.5,color=POSTURE[pp],fontweight="bold")
    for rj,rr in enumerate(roles):
        y=nr-1-rj
        ax.text(-0.12,y+0.5,rolelab[rr],ha="right",va="center",fontsize=8.5,color=ROLE[rr],fontweight="bold")
    ax.text(nc/2,nr+0.62,"Coordination posture  (whose rules you follow)",ha="center",va="bottom",fontsize=8,color=MUTEDHEX)
    ax.text(-1.65,nr/2,"Standards role\n(what you do with the standard)",ha="center",va="center",fontsize=8,color=MUTEDHEX,rotation=90)
    ax.set_xlim(-1.95,nc+0.05); ax.set_ylim(-0.05,nr+0.95); ax.axis("off")
    plt.tight_layout(); path=f"{FIGD}/two_axis_infographic.png"; plt.savefig(path,facecolor=PAPER,bbox_inches="tight"); plt.close()
    return path
MUTEDHEX="#5C5C5C"
INFOG=make_infographic()

# ---------- evaluative scores (same rules as the methodology note) ----------
def lines(s): return [l.strip() for l in (s or "").split("\n") if l.strip()]
def has_year(s): return sum(1 for l in lines(s) if re.search(r"20\d\d", l))
def score(p):
    tl=p.get("migrationTimeline") or ""; ms=has_year(tl); target=bool(p.get("targetCompletion"))
    regs=lines(p.get("mainRegulation")); nreg=len(regs)
    legal=p.get("legalStatus"); binding=legal=="binding"; soft=legal=="soft-only"
    role=p.get("standardsRole"); dom=p.get("dominantProcess"); gov=lines(p.get("govActors")); ngov=len(gov)
    part=bool((p.get("processParticipation") or "").strip())
    hybrid=(p.get("hybridDeployment") or "").strip(); has_hybrid=hybrid not in ("","None stated","None")
    fams=bool((p.get("standardFamilies") or p.get("algorithms") or "").strip())
    text=" ".join([p.get("obligation") or "",p.get("mainRegulation") or "",p.get("summary") or "",tl]).lower()
    eu=("nis2" in text or "dora" in text or "(eu)" in text)
    s={"R1":2 if ms>=1 else (1 if (nreg>0 or target) else 0),"R2":2 if binding else (1 if (soft or dom) else 0),
       "C1":2 if (eu or nreg>=2) else (1 if (nreg==1 or part) else 0),"C2":2 if ngov>=3 else (1 if (ngov>=1 or part) else 0),
       "E1":2 if ms>=1 else (1 if target else 0),"E2":2 if ("inventory" in text or "cbom" in text) else (1 if binding else 0),
       "E3":2 if any(w in text for w in ("monitor","indicator","audit")) else None,"Ef1":2 if binding else (1 if soft else 0),
       "G1":2 if ngov>=2 else (1 if ngov==1 else 0),"G2":2 if any(w in text for w in ("skill","training","academ","programme","workforce")) else None,
       "I1":2 if (role or (dom and (has_hybrid or fams or nreg>0))) else (1 if dom else 0)}
    def mean(k):
        v=[s[x] for x in k if s.get(x) is not None]; return round(sum(v)/len(v),1) if v else None
    return {"Relevance":mean(["R1","R2"]),"Coherence":mean(["C1","C2"]),"Effectiveness":mean(["E1","E2","E3"]),
            "Efficiency":mean(["Ef1"]),"Governance":mean(["G1","G2"]),"Impact":mean(["I1"])}
CONF={"Complete":"High","Partial":"Medium","Placeholder":"Low"}
rows=[]
for p in profs:
    ax=score(p); conf=p.get("confidence") or CONF.get(p.get("dataStatus"),"-")
    rows.append((p.get("country") or p.get("iso3"), ax, conf))
rows.sort(key=lambda r:r[0])

# ---------- document scaffold ----------
doc=Document()
n=doc.styles["Normal"]; n.font.name="Georgia"; n.font.size=Pt(10.5); n.font.color.rgb=INK
n.paragraph_format.space_after=Pt(7); n.paragraph_format.line_spacing=1.22
for hid,sz in (("Heading 1",15),("Heading 2",12),("Heading 3",10.5)):
    st=doc.styles[hid]; st.font.name="Arial"; st.font.size=Pt(sz); st.font.bold=True; st.font.color.rgb=INK
    st.paragraph_format.space_before=Pt(12); st.paragraph_format.space_after=Pt(4); st.paragraph_format.keep_with_next=True
sec=doc.sections[0]; sec.page_width=Cm(21.0); sec.page_height=Cm(29.7)
for m in ("top_margin","bottom_margin","left_margin","right_margin"): setattr(sec,m,Cm(2.3))

def body(t,after=7,size=10.5):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(size); p.paragraph_format.space_after=Pt(after); return p
def h1(t): doc.add_heading(t,1)
def h2(t): doc.add_heading(t,2)
def h3(t): doc.add_heading(t,3)
def bullet(t,size=10.5):
    p=doc.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); p.paragraph_format.space_after=Pt(3)
    for rr in p.runs: rr.font.name="Georgia"
    return p
def caption(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.italic=True; r.font.size=Pt(9); r.font.color.rgb=MUTED; p.paragraph_format.space_after=Pt(10); return p
def shade(cell,fill):
    sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),fill); cell._tc.get_or_add_tcPr().append(sh)
def promptbox(title, body_text):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.rows[0].cells[0]; c.width=Cm(16.4); shade(c,CODEBG)
    tcPr=c._tc.get_or_add_tcPr(); borders=OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        e=OxmlElement(f'w:{edge}'); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'6'); e.set(qn('w:color'),'BBB6AC'); borders.append(e)
    tcPr.append(borders)
    ph=c.paragraphs[0]; ph.paragraph_format.space_after=Pt(3)
    rh=ph.add_run(title); rh.bold=True; rh.font.name="Arial"; rh.font.size=Pt(8.5); rh.font.color.rgb=MUTED
    for ln in body_text.split("\n"):
        pp=c.add_paragraph(); pp.paragraph_format.space_after=Pt(0); pp.paragraph_format.line_spacing=1.16
        rr=pp.add_run(ln); rr.font.name="Consolas"; rr.font.size=Pt(8.6); rr.font.color.rgb=INK
    doc.add_paragraph().paragraph_format.space_after=Pt(3)
def table(headers,data,widths,sizes=9.0,headfill=WASH):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    lay=OxmlElement('w:tblLayout'); lay.set(qn('w:type'),'fixed'); t._tbl.tblPr.append(lay)
    def cell(c,txt,w,bold=False,fill=None):
        c.width=Cm(w); pr=c.paragraphs[0]; pr.paragraph_format.space_after=Pt(1); pr.paragraph_format.line_spacing=1.0
        rr=pr.add_run(txt); rr.bold=bold; rr.font.name="Arial"; rr.font.size=Pt(sizes); rr.font.color.rgb=INK
        if fill: shade(c,fill)
    for i,h in enumerate(headers): cell(t.rows[0].cells[i],h,widths[i],bold=True,fill=headfill)
    for row in data:
        cs=t.add_row().cells
        for i,v in enumerate(row): cell(cs[i],v,widths[i])
    doc.add_paragraph().paragraph_format.space_after=Pt(3); return t

# ===================== TITLE =====================
p=doc.add_paragraph(); r=p.add_run("The QSC Atlas, rebuilt as a scrolling narrative")
r.font.name="Georgia"; r.font.size=Pt(18); r.bold=True; p.paragraph_format.space_after=Pt(2)
p=doc.add_paragraph(); r=p.add_run("A precise UI and UX build brief for Claude Code, grounded in the ui-ux-pro-max design skill and merged with the evaluative-layer methodology")
r.font.name="Arial"; r.font.size=Pt(11.5); r.font.color.rgb=MUTED; p.paragraph_format.space_after=Pt(2)
p=doc.add_paragraph(); r=p.add_run("Working brief, June 2026  |  Site title: The geopolitics of the post-quantum transition  |  Stack: Astro static site, React islands, d3 globe")
r.font.name="Arial"; r.font.size=Pt(9); r.font.color.rgb=MUTED; p.paragraph_format.space_after=Pt(12)

ab=doc.add_paragraph(); abr=ab.add_run("How to use this brief. "); abr.bold=True
ab.add_run("This document is written to be handed to Claude Code (or any capable coding agent) to rebuild the QSC "
  "Atlas homepage as one continuous, scroll-driven story in which a single globe is the through-line. It keeps the "
  "existing stack and the monochrome discipline, and it draws its rules from the ui-ux-pro-max design skill: token "
  "architecture, interaction states, the scroll-storytelling pattern, and the accessibility, performance and motion "
  "guidelines. Parts 0 to 7 are the build. Part 8 carries the methodology the site now presents, including the "
  "evaluative layer and the credibility instrument, so the brief is self-contained. Every instruction the agent "
  "should follow literally is set in a bordered monospaced block.")
for r in ab.runs: r.font.size=Pt(10)

# ===================== PART 0 =====================
h1("Part 0  Orientation")
h2("0.1  What we are building, and why")
body("The homepage today is a globe beside a block of text, then four stacked bands. The rebuild turns it into a "
  "single vertical journey. As the visitor scrolls, the globe stays on stage and changes its place and size while "
  "short columns of text arrive beside it; near the end the globe leaves and the people behind the work appear. The "
  "point is not spectacle. It is that one object, the map of the world, carries three questions in sequence: why this "
  "matters to you, how the Atlas reads the world, and who made it. The scroll is the argument.")
body("This pattern is the design skill's Scroll-Triggered Storytelling pattern (intro hook, then chapters, then a "
  "closing): narrative raises time on page, but only if it respects the reader. So the globe is pinned and "
  "transformed by native scroll, never by hijacking the scrollbar, and the whole thing degrades to plain stacked "
  "sections when motion is not wanted.")
h2("0.2  What to keep, and the three hard rules")
bullet("Keep the stack: Astro static pages, React islands, the existing d3 globe (HeroGlobe), content loaded from "
  "Notion through content.config.ts, and the design tokens already defined in the global stylesheet.")
bullet("Keep the discipline: ink on warm paper, generous whitespace, oversized serif display type, no rounded "
  "corners, no shadows. Colour is meaningful, never decorative: the only colour on screen is a country's bloc "
  "colour. This matches the design skill's Minimalist Monochrome and Exaggerated Minimalism styles.")
bullet("Rule 1, motion is earned: animate at most one or two elements per view; never animate width or height "
  "(transform and opacity only); every effect has a prefers-reduced-motion fallback.")
bullet("Rule 2, never hijack scroll: the page scrolls natively. The globe is pinned with position: sticky and driven "
  "by scroll progress. No scroll-jacking, no forced parallax (the skill flags both as motion-sickness risks).")
bullet("Rule 3, accessibility is not optional: WCAG AA contrast, visible focus rings, full keyboard path, semantic "
  "headings, alt and aria, and no meaning carried by colour alone.")
h2("0.3  The title")
body("Set the site title and hero headline to The geopolitics of the post-quantum transition. Alternatives the team "
  "considered, kept here for the record: The geopolitics of PQC transition (shorter, uses the abbreviation); "
  "Governed everywhere, coordinated nowhere (leads with the thesis); Who governs the post-quantum transition? (a "
  "question). The chosen title is geopolitics-forward and reads to a non-specialist.")

# ===================== PART 1 =====================
h1("Part 1  Design foundations")
h2("1.1  Token architecture")
body("Work in the three-layer token model the design skill prescribes: primitive values, then semantic aliases, then "
  "per-component tokens, all as CSS custom properties. The Atlas already has the primitives and most semantics; this "
  "rebuild adds a small motion and scroll layer. Do not introduce hardcoded colours, durations or spacing in "
  "components; read tokens.")
promptbox("Add to the global token layer (:root)",
  "/* motion (primitive) */\n"
  "--dur-fast: 150ms;  --dur-mid: 240ms;  --dur-slow: 1100ms;\n"
  "--ease-standard: cubic-bezier(.2,.0,.2,1);\n"
  "--ease-out: cubic-bezier(.16,1,.3,1);\n"
  "/* scroll scenes (semantic) */\n"
  "--scene-min-h: 100svh;        /* each scene is at least one viewport */\n"
  "--globe-stage-w: clamp(280px, 42vw, 620px);\n"
  "--rail-w: 2px;                /* progress rail */\n"
  "/* display type (semantic) */\n"
  "--type-display: clamp(2.2rem, 1.2rem + 3.4vw, 4.6rem);  /* hero H1 */\n"
  "--type-scene:   clamp(1.5rem, 1.0rem + 1.6vw, 2.4rem);  /* scene H2 */")
body("Keep the existing palette tokens. For reference, the bloc colours are the only accents allowed on the page: "
  "EU 5b54a8, NIST-bloc 2b4c7e, sovereign-bloc 7a3b5e, engaged-unaligned 6b7280; the standards-role accents "
  "(maker a8322a, contextualiser b9762e, taker 4a6fa5, sovereign-developer 7a3b5e) appear only where the role is the "
  "subject. Ink is 1A1A1A, paper F7F5F0, hairline D8D6D0.")
h2("1.2  Type, space and the editorial style")
body("Display headings use the serif (Newsreader, var(--font-reading)) at the oversized display size above; running "
  "text uses the instrument sans (Schibsted Grotesk); data, counts and labels use the mono (Spline Sans Mono). "
  "Whitespace does the framing: at least one viewport per scene, wide outer margins, a single column of text no "
  "wider than the existing --measure. No card chrome, no borders except the hairline and the progress rail.")
h2("1.3  Motion specification")
body("Transitions follow the skill's table: colour 150ms ease-in-out, transform 200ms ease-out, opacity 150ms. "
  "Scene entrances use a short fade-and-rise (translateY 16px to 0, opacity 0 to 1) over --dur-mid, staggered by no "
  "more than two elements. The globe's own transform across scenes is eased with --ease-out and is the one piece of "
  "sustained motion on the page. Under prefers-reduced-motion: reduce, all entrances resolve instantly and the globe "
  "is placed in its final per-scene position without tweening.")
h2("1.4  Accessibility and performance budget")
bullet("Contrast: body text on paper is already AA; never drop muted text below the existing --ink-muted on paper.")
bullet("Keyboard: every link and control reachable in DOM order; visible focus ring (var(--focus-ring)); a skip link "
  "to main content.")
bullet("Headings: one H1 (the hero title); each scene a single H2; supporting pages keep the order.")
bullet("Colour-not-only: the bloc colour is always paired with its text label and, on the globe, with a shape or "
  "position cue, so the legend never relies on hue alone.")
bullet("Performance: reserve the globe's box to avoid layout shift (CLS under 0.1); load the globe island with "
  "client:idle as today; font-display: swap; throttle any scroll listener with requestAnimationFrame; keep "
  "per-frame work under about 16ms.")

h2("1.5  Reference scales")
table(["Type token","Size","Use"],
 [("--type-display","clamp(2.2rem, 1.2rem+3.4vw, 4.6rem)","Hero H1 only"),
  ("--type-scene","clamp(1.5rem, 1.0rem+1.6vw, 2.4rem)","Scene H2"),
  ("--text-lg / base","1.125rem / 1rem","Lede / body"),
  ("--text-sm / xs","0.875rem / 0.78rem","Captions, labels"),
  ("mono (data)","0.72-0.9rem","Counts, scores, ISO codes"),
 ],[3.4,6.6,6.4],sizes=8.6)
caption("Table (type). Display is Newsreader serif; body is Schibsted Grotesk; data is Spline Sans Mono.")
table(["Motion","Duration","Easing"],
 [("Colour and background","150ms","ease-in-out"),
  ("Transform (scene entrance)","200ms","ease-out"),
  ("Opacity","150ms","ease"),
  ("Globe pose across scenes","~1100ms","--ease-out"),
  ("Reduced motion","0ms","none, render final state"),
 ],[5.4,3.4,7.6],sizes=8.6)
caption("Table (motion). From the design skill's transition table; the globe pose is the one sustained motion.")
table(["Accessibility check","Target"],
 [("Contrast","Body AA 4.5:1; large text 3:1"),
  ("Focus","Visible ring on every control (var(--focus-ring))"),
  ("Keyboard","Full path in DOM order; skip link to main"),
  ("Headings","One H1; one H2 per scene; no skipped levels"),
  ("Colour-not-only","Bloc colour always paired with a text label"),
  ("Motion","prefers-reduced-motion honoured everywhere"),
 ],[5.0,11.4],sizes=8.6)
caption("Table (accessibility). The skill's priority-1 checks, applied here.")
table(["Performance budget","Target"],
 [("CLS","Under 0.1 (reserve the globe box)"),
  ("Globe island","client:idle; cap device pixel ratio at 2"),
  ("Scroll work","requestAnimationFrame; observer-driven, not per scroll event"),
  ("Fonts","font-display: swap; preload only the display face"),
  ("Frame budget","Under 16ms; pause globe animation when offscreen"),
 ],[5.0,11.4],sizes=8.6)
caption("Table (performance). The skill's priority-3 checks.")

# ===================== PART 2 =====================
h1("Part 2  The scroll architecture")
h2("2.1  Structure")
body("Build the homepage as a single relative container holding two layers. The back layer is the globe stage, a "
  "position: sticky element that stays in view for the length of the narrative. The front layer is a column of scene "
  "panels stacked vertically, each at least one viewport tall, each holding the text for its scene. As the panels "
  "scroll past, an IntersectionObserver marks the active scene and the globe transforms to that scene's pose. "
  "Because the panels scroll natively and the globe only reads progress, the scrollbar is never captured.")
promptbox("Homepage skeleton (src/pages/index.astro)",
  "<main class=\"scrolly\">\n"
  "  <div class=\"globe-stage\" data-scene=\"hero\">\n"
  "    <HeroGlobe client:idle mapProcess={mapProcess} />\n"
  "  </div>\n"
  "  <section class=\"scene scene--hero\"  data-scene=\"hero\">  ...text left...  </section>\n"
  "  <section class=\"scene scene--method\" data-scene=\"method\"> ...text left, globe right... </section>\n"
  "  <section class=\"scene scene--team\"  data-scene=\"team\">  ...team, globe gone...  </section>\n"
  "  <nav class=\"scroll-rail\" aria-hidden=\"true\">...three dots...</nav>\n"
  "</main>")
promptbox("Pinning and the scene driver (concept)",
  ".scrolly { position: relative; }\n"
  ".globe-stage { position: sticky; top: 0; height: 100svh;\n"
  "  display: grid; place-items: center; transition: transform var(--dur-slow) var(--ease-out),\n"
  "  opacity var(--dur-mid) var(--ease-standard); will-change: transform, opacity; }\n"
  ".scene { min-height: var(--scene-min-h); display: grid; align-items: center; }\n"
  "/* JS: IntersectionObserver(threshold .5) on .scene -> set body[data-scene];\n"
  "   CSS maps [data-scene] to a .globe-stage transform (see 2.2-2.4). */")
promptbox("Scene driver (src/scripts/scenes.ts, loaded once)",
  "const reduce = matchMedia('(prefers-reduced-motion: reduce)').matches;\n"
  "const small  = matchMedia('(max-width: 860px)').matches;\n"
  "if (!reduce && !small) {\n"
  "  const io = new IntersectionObserver((entries) => {\n"
  "    for (const e of entries) if (e.isIntersecting)\n"
  "      document.body.dataset.scene = e.target.getAttribute('data-scene');\n"
  "  }, { threshold: 0.5 });\n"
  "  document.querySelectorAll('.scene').forEach((s) => io.observe(s));\n"
  "}  /* CSS reads body[data-scene] and sets the .globe-stage transform */")
promptbox("Scene text column (text always left)",
  ".scene { grid-template-columns: minmax(0, 38rem) 1fr; gap: var(--space-8);\n"
  "  max-width: var(--content-wide); margin-inline: auto; padding: var(--space-6); }\n"
  ".scene .text { grid-column: 1; }              /* text column on the left */\n"
  ".scene--team { grid-template-columns: 1fr; }  /* globe gone: single column */\n"
  "@media (max-width: 860px) { .scene { grid-template-columns: 1fr; } }")
h2("2.2  Scene 1, the hero: the globe comes forward")
body("At the top the globe is the largest it will be and sits to the right of centre; the text column is on the left. "
  "As scene 1 settles, the globe scales from about 0.9 to 1.0 and drifts very slightly forward (a small scale, not a "
  "zoom). The left column carries the H1 (the new title), one short answer to why the Atlas is useful to the reader, "
  "a link to the orientation page, and the thesis in a few words. Copy is in Part 3.")
promptbox("Scene 1 globe pose",
  "body[data-scene=\"hero\"] .globe-stage { transform: translateX(8%) scale(1.0); opacity: 1; }")
h2("2.3  Scene 2, the methodology: the globe moves right")
body("On scroll into scene 2 the globe translates further right and shrinks to make room, settling as a quiet figure "
  "beside the text. The left column explains how the Atlas reads the world: the two axes, what each category means, "
  "and why it matters, with a compact infographic and a link to the full methodology page. The infographic is the "
  "two-axis grid in 3.2; build it as inline SVG so it inherits the tokens.")
promptbox("Scene 2 globe pose",
  "body[data-scene=\"method\"] .globe-stage { transform: translateX(34%) scale(0.62); opacity: .9; }")
h2("2.4  Scene 3, the team: the globe leaves")
body("Entering scene 3 the globe fades and scales down out of the stage, returning the page to plain ink on paper for "
  "the human close. The team section names the two leads with their CEPS roles and links, and links to the Task "
  "Force's two reports. After the globe leaves, restore normal document flow (the sticky stage ends) so the footer "
  "sits naturally.")
promptbox("Scene 3 globe pose, then release",
  "body[data-scene=\"team\"] .globe-stage { transform: translateX(34%) scale(0.2); opacity: 0; }\n"
  "/* end the sticky stage at the team section so the footer flows normally */")
h2("2.5  The progress rail")
body("A thin fixed rail of three marks (hero, methodology, team) sits at the right edge, the active mark filled in "
  "ink, the rest hairline. It is decorative for sighted orientation only (aria-hidden), because the real navigation "
  "is the native scroll and the header links. Clicking a mark smooth-scrolls to that scene unless reduced motion is "
  "set, in which case it jumps.")
h2("2.6  Mobile and reduced motion")
body("Below about 860px, and whenever prefers-reduced-motion: reduce is set, drop the pinned stage entirely. Each "
  "scene becomes a normal stacked block: a static globe figure (or a still image of it) above its text, in source "
  "order. No transforms, no observers driving motion. This is the design skill's instruction to simplify animations "
  "on mobile and to honour motion preferences, and it keeps the content identical for everyone.")

# ===================== PART 3 =====================
h1("Part 3  Scene content and copy")
body("Use this copy directly. It is concise by design: the reader should grasp each scene in one breath, then follow "
  "a link to go deeper.")
h2("3.1  Scene 1 copy")
body("H1: The geopolitics of the post-quantum transition", after=3)
body("Lede (why it is useful to you): See, for any country, what moving to quantum-safe encryption actually requires: "
  "the rule that binds, the deadline that applies, and how credible the plan is. Built only from official sources, in "
  "one place.", after=3)
body("Link: Who is this for, and how to read the map ->  (to the orientation page).", after=3)
body("Thesis, in a few words (the detailed version): Post-quantum cryptography is governed transnationally. No single "
  "state or body sets the rules, yet the standards that matter, the NIST algorithms, are adopted almost everywhere, "
  "so authority is dispersed across standards bodies, regulators and national agencies at once. Coordination stops at "
  "the algorithm. Around that shared core the transition fragments in two ways the Atlas makes visible. Regulatory "
  "fragmentation: the same algorithms are wrapped in very different instruments, binding EU duties under NIS2 and "
  "DORA, soft national roadmaps elsewhere, and nothing at all across much of the world, so what an organisation must "
  "do depends entirely on where it operates. Roadmap fragmentation: the clocks diverge, the EU's coordinated 2030 and "
  "2035 milestones, France's phased not-earlier-than schedule, the United States' executive deadlines, China's "
  "separate SM-series track, so the same migration runs to different timetables in different places. One quiet, shared "
  "technical core; many divergent governance wrappers. That gap is the subject of the Atlas.")
h2("3.2  Scene 2 copy and the infographic")
body("Heading: How the Atlas reads the world. Lead line: every country is placed on two axes, then read for what it "
  "must do and how credible its plan is.")
try:
    doc.add_picture(INFOG, width=Cm(15.6))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    body(f"[infographic: {e}]")
caption("Sample of the scene-2 infographic (build as inline SVG on the site): the two axes, with the live count of "
  "countries in each cell. Columns are coordination posture; rows are standards role.")
body("Category copy (what each means, and why it matters):", after=3)
bullet("Coordination posture, whose rules you follow: EU, NIST-bloc, sovereign-bloc, or engaged but unaligned. Why it "
  "matters: it tells you whose timeline and whose obligations apply to you.")
bullet("Standards role, what a country does with the standard: standard-maker, contextualiser, standard-taker, or "
  "sovereign-developer. Why it matters: it shows who shapes the rules and who only adopts them, which is where power "
  "sits.")
bullet("Regulatory basis, the instrument behind the posture: the main regulation, whether it is binding, and the "
  "plain obligation. Why it matters: it turns a label into what you actually have to do.")
bullet("Governance credibility, the evaluative layer: six axes scored from the sources. Why it matters: it tells you "
  "how credible a plan is, not merely that one exists.")
body("Link: Read the full methodology, with the scoring and the sources ->  (to the methodology page).")
h2("3.3  Scene 3 copy, the team")
body("Heading: Who is behind the Atlas. The Atlas is maintained within the Cybersecurity@CEPS Initiative at the "
  "Centre for European Policy Studies.", after=3)
bullet("Swann Ashworth, Associate Research Assistant, GRID Unit, Cybersecurity@CEPS; rapporteur for the CEPS Task "
  "Force on the EU transition to a quantum-safe world. ceps.eu/ceps-staff/swann-ashworth/")
bullet("Lorenzo Pupillo, Associate Senior Research Fellow and Head of the Cybersecurity@CEPS Initiative. "
  "ceps.eu/ceps-staff/lorenzo-pupillo/")
body("From the Task Force (link both):", after=3)
bullet("Quantum Technologies and Cybersecurity: technology, governance and policy challenges (CEPS, 2023). "
  "ceps.eu/ceps-publications/quantum-technologies-and-cybersecurity/")
bullet("Strengthening the EU transition to a quantum-safe world: technology, market, governance and policy challenges "
  "(CEPS Task Force report, 2025). ceps.eu/ceps-publications/strengthening-the-eu-transition-to-a-quantum-safe-world/")

# ===================== PART 4 =====================
h1("Part 4  The two supporting pages")
h2("4.1  Orientation page: who is this for, how to navigate")
body("A short page linked from scene 1. Three blocks: who it is for (policymakers, CISOs and advisers, researchers), "
  "what they get (per-country obligations, deadlines, credibility, all sourced), and how to read it (the two axes, "
  "the colour rule, the confidence grade, how to open a country). Same monochrome layout, no globe. End with a link "
  "into the map.")
h2("4.2  Methodology page: the full model")
body("The complete version of scene 2, linked from there. It carries the two-axis classification model, the "
  "regulatory layer, and the evaluative layer in full, with the scoring scheme, the verification rules, and the "
  "sources. Reuse the content in Part 8 of this brief. Build it as a long reading page with a sticky in-page table of "
  "contents on the left, the same infographic, and the credibility instrument shown live for one example country.")

# ===================== PART 5 =====================
h1("Part 5  The credibility instrument (country profile)")
body("This is the component the methodology produces, specified for build. It belongs on each country profile, below "
  "the Regulatory basis block, under a heading Governance credibility, and it is the live element referenced on the "
  "methodology page.")
h2("5.1  Data shape")
body("Add an optional evaluation object to each profile (data/profiles/<ISO3>.json), mirror it in Notion, and type it "
  "in content.config.ts as nullable so profiles without it render an empty gauge. evaluation maps each of the six "
  "criteria (relevance, coherence, effectiveness, efficiency, governance, impact) to { score (0 to 2 or null), "
  "confidence (High|Medium|Low), subScores: [{ key, label, score (0|1|2|null), reason, sourceUrl }] }.")
promptbox("evaluation field, example (data/profiles/FRA.json)",
  "\"evaluation\": {\n"
  "  \"relevance\": { \"score\": 2.0, \"confidence\": \"High\", \"subScores\": [\n"
  "    { \"key\": \"R1\", \"label\": \"Transition plan\", \"score\": 2,\n"
  "      \"reason\": \"ANSSI phased roadmap published\", \"sourceUrl\": \"https://...\" },\n"
  "    { \"key\": \"R2\", \"label\": \"Threat articulation\", \"score\": 2,\n"
  "      \"reason\": \"Binding EU regime\", \"sourceUrl\": \"https://...\" } ] },\n"
  "  \"effectiveness\": { \"score\": 1.5, \"confidence\": \"High\", \"subScores\": [ ... ] }\n"
  "  /* coherence, efficiency, governance, impact follow the same shape */\n"
  "}")
promptbox("Schema wiring (src/content.config.ts, zod, nullable)",
  "const subScore = z.object({ key: z.string(), label: z.string(),\n"
  "  score: z.number().nullable(), reason: z.string(), sourceUrl: z.string().url().optional() });\n"
  "const axis = z.object({ score: z.number().nullable(),\n"
  "  confidence: z.enum(['High','Medium','Low']), subScores: z.array(subScore) });\n"
  "evaluation: z.object({ relevance: axis, coherence: axis, effectiveness: axis,\n"
  "  efficiency: axis, governance: axis, impact: axis }).nullable().optional(),")
h2("5.2  Component, interaction, accessibility")
bullet("Render src/components/CredibilityGauge.tsx as a React island: an SVG six-spoke radial gauge (viewBox 0 0 240 "
  "240), one axis per criterion, each filled to score/2 of the radius, the six points joined into a polygon filled "
  "in the country's bloc colour at low alpha, the bloc colour the only accent on the outer ring.")
bullet("Labels in Schibsted Grotesk outside each spoke, the numeric score in Spline Sans Mono. A null axis renders as "
  "a dashed grey spoke labelled not yet assessed; a low-confidence axis renders at reduced opacity.")
bullet("On mount the gauge assembles (spokes draw, then points scale in) over about 1.1s eased with --ease-out; "
  "hovering or focusing a spoke opens a small panel with that criterion's sub-scores, each with its value, one-line "
  "reason, a source link and the confidence. A compare control overlays a second country as a ghosted dashed outline.")
bullet("Accessibility: svg role=img with title and desc; each spoke a focusable button with an aria-label of the form "
  "Relevance: 2.0 of 2, established; values shown as text, never colour alone; under reduced motion the gauge renders "
  "in its final state with no draw-on and the panel shown inline.")
body("The scoring behind the gauge, the six criteria and their verification, is in Part 8. The first-pass scores for "
  "every country are in Table 3, ready to seed the evaluation field.")

# ===================== PART 6 =====================
h1("Part 6  Components and states")
body("Every interactive element carries the standard state set from the design skill, in priority order disabled, "
  "loading, active, focus, hover, default, with the transitions from 1.3. The new or changed components:")
table(["Component","Role","Key states / notes"],
 [("Globe stage","Pinned through-line; transforms per scene","Three poses (2.2-2.4); reduced-motion = static per scene"),
  ("Scene panel","Holds one scene's text","Entrance fade-rise on first view; one H2 each"),
  ("Scroll rail","Orientation marks (3)","Active mark ink, rest hairline; aria-hidden; click = scroll/jump"),
  ("Infographic","Two-axis grid (scene 2)","Inline SVG from tokens; cells labelled, not colour-only"),
  ("Credibility gauge","Six-axis score (profile)","Hover/focus reveal; compare overlay; na and low-confidence states"),
  ("Team card","Person name, role, link","Hover underline; focus ring; external links rel=noopener"),
  ("Primary link","Scene CTAs to pages","One primary per scene; underline on hover; 44px target"),
 ],[3.0,5.0,8.4],sizes=8.6)
caption("Table 1. New and changed components, with their states.")

# ===================== PART 7 =====================
h1("Part 7  Implementation plan for Claude Code")
body("Work in this order; commit after each step; run the build and the checks before moving on.")
def step(nn,t):
    p=doc.add_paragraph(); r=p.add_run(f"Step {nn}.  "); r.bold=True; r.font.name="Arial"; r.font.size=Pt(10)
    r2=p.add_run(t); r2.font.size=Pt(10.5); p.paragraph_format.space_after=Pt(4)
step(1,"Tokens. Add the motion, scroll and display-type tokens from 1.1 to the global stylesheet. No visual change yet.")
step(2,"Scaffolding. Rebuild src/pages/index.astro to the skeleton in 2.1: a .scrolly container, the sticky "
  ".globe-stage wrapping the existing HeroGlobe, three .scene sections, and the scroll rail. Keep the header and "
  "footer. Move the old bands (Documents, About, Advisory) to the footer or the orientation page; they are not part "
  "of the narrative.")
step(3,"Scene driver. Add a tiny client script: an IntersectionObserver over .scene that sets body[data-scene]; map "
  "each value to a .globe-stage transform in CSS (2.2 to 2.4). Throttle nothing here, the observer is event-driven. "
  "Guard everything behind a prefers-reduced-motion and a min-width check.")
step(4,"Copy. Drop in the Part 3 copy. One H1, one H2 per scene. Wire the two links to the pages built in step 6.")
step(5,"Infographic. Build the two-axis grid (3.2) as an inline SVG component reading the live country counts from the "
  "collection, coloured only by bloc, every cell labelled.")
step(6,"Pages. Build the orientation page (4.1) and the methodology page (4.2) from Part 8. Add the routes to the nav.")
step(7,"Credibility gauge. Implement CredibilityGauge.tsx (Part 5), add the evaluation field to the schema and to one "
  "example profile, and mount the gauge on the profile and on the methodology page.")
step(8,"Reduced motion and mobile. Verify the stacked fallback (2.6): no observers driving motion, static globe, "
  "source-order content.")
step(9,"QA and build. Run the acceptance checks below, then astro build.")
h2("7.1  Acceptance criteria")
bullet("The page scrolls natively; the globe is pinned and moves through three poses; the scrollbar is never captured.")
bullet("Exactly one H1; each scene one H2; tab order matches reading order; focus rings visible throughout.")
bullet("Under prefers-reduced-motion the globe does not tween and entrances are instant; on mobile the layout is "
  "stacked with a static globe.")
bullet("The only colours rendered are bloc colours; every coloured element has a text label; AA contrast holds.")
bullet("No layout shift on load (globe box reserved); Lighthouse performance and accessibility both green; astro "
  "build passes.")
bullet("Team links and Task Force links resolve to the CEPS URLs in 3.3; the methodology page shows a live gauge.")

# ===================== PART 8  GLOBE =====================
h1("Part 8  Globe integration")
body("The globe is the one object that persists across the narrative, so its integration is the load-bearing piece "
  "of the build.")
h2("8.1  What the globe is, and what changes")
body("HeroGlobe is the existing React island that renders a slowly rotating globe coloured by coordination posture, "
  "keyed by country, with the posture map passed in as mapProcess. Three changes make it fit the scroll stage. "
  "First, it must fill its container and follow the container's size, so wrap it in the sticky stage and let a CSS "
  "transform scale the stage rather than re-rendering the globe; add a ResizeObserver so the canvas keeps its "
  "resolution when the box changes. Second, it keeps its idle rotation in scene 1 but pauses under reduced motion. "
  "Third, it stays clickable in scene 1, where a country opens its profile, and becomes non-interactive once it "
  "shrinks in scenes 2 and 3.")
promptbox("Globe stage wrapper (props)",
  "<div class=\"globe-stage\" data-scene={scene}>\n"
  "  <HeroGlobe client:idle mapProcess={mapProcess}\n"
  "    interactive={scene === 'hero'}   /* clicks open profiles only in scene 1 */\n"
  "    autoRotate={!reducedMotion}       /* still globe under reduced motion */\n"
  "    fit=\"container\" />               /* size to the stage via ResizeObserver */\n"
  "</div>")
h2("8.2  Reduced motion and mobile")
body("Under prefers-reduced-motion the globe renders as a still sphere in each scene's final pose, with no rotation "
  "and no tweening between scenes. On small screens the globe becomes a static figure above each scene's text in "
  "source order; if the live canvas is costly on low-end devices, ship a pre-rendered still as the mobile fallback.")

# ===================== PART 9  IA + COPY =====================
h1("Part 9  Information architecture, navigation, and page copy")
h2("9.1  Routes and navigation")
bullet("Routes: / (the narrative), /who (orientation), /methodology, /map (the flat map), /countries and "
  "/countries/<iso3>, /documents.")
bullet("Header: Map, Countries, Methodology, Documents, About. Footer: a short Documents summary, the Advisory line, "
  "an About and methodology link, contact, and the sources-only note.")
bullet("The old homepage bands (Documents, About, Advisory) move to the footer and their own pages; they are not "
  "scenes in the narrative.")
h2("9.2  Orientation page copy (/who)")
body("Heading: Who the Atlas is for, and how to read it.", after=3)
bullet("Who it is for: policymakers comparing national positions; CISOs and advisers scoping a transition across "
  "jurisdictions; researchers who need sourced, comparable data.")
bullet("What you get: for any country, the binding instrument, the legal status, the plain obligation, the deadline, "
  "and a credibility reading, every claim linked to an official source.")
bullet("How to read it: colour shows only the standards process; the two axes are coordination posture and standards "
  "role; each profile carries a confidence grade and a last-verified date; click a country to open it.")
body("Close with a primary link: Open the map ->.")
h2("9.3  Methodology page structure (/methodology)")
body("A long reading page with a sticky in-page contents on the left. Sections in order: the two axes (with the "
  "scene-2 infographic); the regulatory layer; the evaluative layer (the six criteria and scoring from Part 11); the "
  "verification rules; one live credibility gauge for an example country; and the sources. Same monochrome layout, "
  "no globe.")

# ===================== PART 10  QA + RATIONALE =====================
h1("Part 10  Quality assurance and design rationale")
h2("10.1  Pre-ship checklist")
table(["Priority (skill)","Check in this build"],
 [("1 Accessibility","AA contrast; focus rings; keyboard path; one H1; colour-not-only; skip link"),
  ("2 Interaction","44px targets; no hover-only actions; clicks open profiles"),
  ("3 Performance","CLS under 0.1; client:idle globe; rAF scroll; font-display swap"),
  ("4 Style","Monochrome; bloc colour only; no radius, no shadow; SVG not emoji"),
  ("5 Layout","Native scroll; no horizontal scroll; mobile stacked fallback"),
  ("6 Type and colour","16px base; serif display; semantic tokens, no raw hex in components"),
  ("7 Animation","1-2 elements per view; transform and opacity only; reduced-motion fallback"),
  ("9 Navigation","Header links; rail is orientation only; browser back intact"),
 ],[3.8,12.6],sizes=8.4)
caption("Table 4. The design skill's priority checks, mapped to this build.")
h2("10.2  Why these choices")
table(["Decision","Grounding in the design skill"],
 [("Pinned globe, native scroll","Scroll-Triggered Storytelling pattern; scroll-jacking and parallax flagged as motion-sickness risks"),
  ("Monochrome, one accent","Minimalist Monochrome and Exaggerated Minimalism styles"),
  ("Three-layer tokens","Token architecture, primitive to semantic to component"),
  ("One or two animations per view","Animation guideline: excessive motion distracts"),
  ("Reduced-motion fallback","Accessibility: respect prefers-reduced-motion"),
  ("Colour paired with labels","Accessibility: do not convey meaning by colour alone"),
 ],[5.2,11.2],sizes=8.4)
caption("Table 5. Each decision and the guideline it follows.")
h2("10.3  Reading the credibility instrument: France")
body("A worked example for the methodology page. France scores high on relevance, coherence, efficiency, governance "
  "and impact, and a step lower on effectiveness. The reading is legible from the axes: ANSSI publishes a phased "
  "transition roadmap and France adopts the NIST baseline with mandatory hybridisation (relevance and impact), under "
  "binding EU instruments with a mandated lead (coherence, efficiency, governance); effectiveness sits lower because "
  "the milestones are not-earlier-than windows and no dedicated measurement duty is recorded. Each axis on the gauge "
  "links to the source behind it, so the reading is auditable rather than asserted.")

# ===================== PART 11  METHODOLOGY =====================
h1("Part 11  The methodology the site presents (merged)")
body("This is the content for the methodology page and the basis of the credibility gauge. It is the evaluative layer "
  "in brief, with the verification rules and the first-pass scores.")
h2("8.1  Six criteria, eleven sub-criteria")
body("The layer adapts Paglieri et al.'s (2025) six-criteria comparative method for national quantum strategies, "
  "sharpened by Vance's (2025) split of evaluation into process, impact and cost-benefit, with two checks drawn from "
  "the CEPS governance recommendations (Pupillo et al., 2025) and the threat criterion from Pupillo et al. (2023). It "
  "sits inside governance of expectations (Budde and Konrad, 2019): a score is a reading at a date, not a verdict.")
table(["Criterion","Sub-criteria (0 / 1 / 2 each)"],
 [("Relevance","R1 transition plan; R2 threat articulation"),
  ("Coherence","C1 policy integration; C2 consultation"),
  ("Effectiveness","E1 milestones; E2 inventory duty; E3 measurement"),
  ("Efficiency","Ef1 proportionate instrument mix"),
  ("Governance quality","G1 mandate and authority; G2 capacity and skills"),
  ("Impact (structured-for)","I1 exposure-reducing adoption"),
 ],[4.2,12.2],sizes=9.0)
caption("Table 2. The six criteria as eleven sub-criteria, each scored 0 absent, 1 partial, 2 present. Reported as a "
  "six-axis profile, not a single number.")
h2("8.2  Verification rules")
bullet("No sub-score is recorded unless it traces to a sourced instrument; an unsupported axis reads as not assessed, "
  "never a default, and the axis is the mean of the sub-criteria that could be assessed.")
bullet("A second reviewer scores each country independently and the two reconcile before a profile is shown.")
bullet("The Atlas confidence grade is carried onto each axis; a country on thin evidence reads as tentative.")
h2("8.3  First-pass scores for every country")
body("The table is a first pass, produced by applying the rules mechanically to the fields each profile records "
  "(timeline and target for R1 and E1; binding status for R2, E2 and Ef1; EU or multiple instruments for C1; the "
  "breadth of recorded bodies for C2 and G1; the regulation and obligation text for E2, E3 and G2; the standards "
  "role with the hybrid stance for I1). A dash means an axis could not be assessed from the recorded fields. Reviewer "
  "reconciliation is pending. These values seed the evaluation field for the gauge.")
def fmt(x): return "-" if x is None else f"{x:.1f}"
data=[(c,fmt(a["Relevance"]),fmt(a["Coherence"]),fmt(a["Effectiveness"]),fmt(a["Efficiency"]),fmt(a["Governance"]),fmt(a["Impact"]),conf) for c,a,conf in rows]
table(["Country","Rel","Coh","Eff","Effi","Gov","Imp","Conf."],data,[3.7,1.55,1.55,1.55,1.55,1.55,1.55,1.6],sizes=8.0)
caption(f"Table 3. First-pass six-axis scores for {len(rows)} countries (0 to 2 per axis). Derived from the recorded "
  "fields; reviewer reconciliation pending.")

# ===================== REFERENCES =====================
h1("References")
refs=[
 "Budde, B., & Konrad, K. (2019). Tentative governing of fuel cell innovation in a dynamic network of expectations. Research Policy, 48(5), 1098-1112.",
 "Paglieri, L., Bonomi Savignon, A., Scalabrini, F., & Costumato, L. (2025). Navigating the quantum frontier: Examining government strategy to the next technological revolution. Transforming Government: People, Process and Policy, 19(4).",
 "Pupillo, L., Ferreira, A., Lipiainen, V., & Polito, C. (2023). Quantum technologies and cybersecurity: Technology, governance and policy challenges. Centre for European Policy Studies. https://www.ceps.eu/ceps-publications/quantum-technologies-and-cybersecurity/",
 "Pupillo, L., et al. (2025). Strengthening the EU transition to a quantum-safe world: Technology, market, governance and policy challenges (Task Force Report). Centre for European Policy Studies. https://www.ceps.eu/ceps-publications/strengthening-the-eu-transition-to-a-quantum-safe-world/",
 "Vance, A. S. (2025). Cybersecurity and quantum computing: A quantitative analysis proposing a framework for assessing quantum cybersecurity maturity [Doctoral dissertation].",
 "nextlevelbuilder. (2025). ui-ux-pro-max design skill (styles, landing patterns, UX guidelines, token architecture, states and variants). https://github.com/nextlevelbuilder/ui-ux-pro-max-skill",
]
for ref in refs:
    p=doc.add_paragraph(); r=p.add_run(ref); r.font.size=Pt(9.5)
    p.paragraph_format.left_indent=Cm(1.0); p.paragraph_format.first_line_indent=Cm(-1.0); p.paragraph_format.space_after=Pt(5)

doc.save(OUT)
print("saved:",OUT,"| countries:",len(rows),"| paragraphs:",len(doc.paragraphs),"| tables:",len(doc.tables))
