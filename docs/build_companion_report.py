#!/usr/bin/env python3
# Builds the QSC Atlas companion report (.docx) from prose + figures + the 54-country
# appendix data. British English, no em-dashes, inline APA. Run in the sandbox.
import json, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ASSETS = "/sessions/gracious-kind-wozniak/mnt/qsc-atlas/docs/report-assets"
OUT = "/sessions/gracious-kind-wozniak/mnt/qsc-atlas/docs/QSC_Atlas_Companion_Report.docx"
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5C, 0x5C, 0x5C)

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Georgia"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.25
for hid, sz in (("Heading 1", 15), ("Heading 2", 12)):
    st = doc.styles[hid]
    st.font.name = "Arial"; st.font.size = Pt(sz); st.font.bold = True
    st.font.color.rgb = INK
    st.paragraph_format.space_before = Pt(14); st.paragraph_format.space_after = Pt(6)

sec = doc.sections[0]
sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Cm(2.4))

def body(text, italic=False, size=None, after=8):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = italic
    if size: r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(after)
    return p

def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)

def figure(fname, caption, width_cm=15.0):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(os.path.join(ASSETS, fname), width=Cm(width_cm))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = c.add_run(caption); cr.italic = True; cr.font.size = Pt(9); cr.font.color.rgb = MUTED
    c.paragraph_format.space_after = Pt(12)

# ---- title block ----
t = doc.add_paragraph(); t.paragraph_format.space_after = Pt(2)
tr = t.add_run("Quantum-safe cryptography: the quiet migration reshaping digital sovereignty")
tr.font.name = "Georgia"; tr.font.size = Pt(19); tr.bold = True; tr.font.color.rgb = INK
s = doc.add_paragraph(); s.paragraph_format.space_after = Pt(2)
sr = s.add_run("A companion report to the QSC Atlas"); sr.font.name = "Arial"; sr.font.size = Pt(12); sr.font.color.rgb = MUTED
m = doc.add_paragraph(); m.paragraph_format.space_after = Pt(14)
mr = m.add_run("Companion to the QSC Atlas (qsc-atlas.vercel.app). Research note, June 2026. [Author and affiliation]")
mr.font.name = "Arial"; mr.font.size = Pt(9); mr.font.color.rgb = MUTED

# ---- abstract ----
ab = doc.add_paragraph(); abr = ab.add_run("Abstract. "); abr.bold = True
ab.add_run("The QSC Atlas maps, country by country, how governments are rebuilding public-key "
  "cryptography against the prospect of a quantum computer able to break it. This report explains "
  "what the Atlas is, the standards-governance principles it rests on, and how to read its categories. "
  "It sets out a two-axis classification: a coordination posture, the collective transition a state "
  "belongs to, and a standards role, the state's relationship to the algorithms themselves. It explains "
  "what it means for a state to have transitioned as a setter, a contextualiser, a taker, or a sovereign "
  "developer, and why the coordination method and standards choice of each country carry theoretical, "
  "methodological, and geopolitical weight. The report documents the method behind the Atlas, including "
  "its use of AI assistance and its design for continuous updating, and classifies the 54 countries "
  "currently recorded.")
for r in ab.runs: r.font.size = Pt(10)

# ---- 1 ----
h1("1. The instrument and its question")
body("Public-key cryptography secures most of what the digital economy runs on, from interbank "
  "settlement to the authentication of government systems. A sufficiently powerful quantum computer "
  "would break the asymmetric algorithms that work depends on. No such machine exists, and the date of "
  "its arrival cannot be known. Governments, regulators, standards bodies and technology firms are "
  "nonetheless already rebuilding the cryptographic foundations of the digital economy against the "
  "anticipated threat. The QSC Atlas records that rebuilding as it happens.")
body("Two features make the transition hard to read from any single vantage point. The first is that "
  "cryptographic protocols are infrastructure in the precise sense Star (1999) gives the term: they "
  "become visible only on breakdown, and they are migrated, not switched. The second is that no single "
  "authority governs the transition. Coordination happens, but it is scattered across national agencies, "
  "regional roadmaps and transnational standards bodies. The Atlas is built to show that scattered "
  "structure rather than to assume a centre that does not exist.")
body("The map answers two questions that are easy to confuse and that matter for different reasons. The "
  "first is which collective transition a country belongs to: whose roadmap, timeline and rules it works "
  "to. The second is what the country does with the algorithms themselves: whether it writes them, adapts "
  "them, adopts them, or builds its own. The first is the map colour; the second is shown on each country "
  "profile. Keeping them apart is the central design decision, and section 3 sets out why.")

# ---- 2 ----
h1("2. Theoretical principles")
body("The Atlas rests on the standards-governance literature, which treats technical standards as "
  "instruments of power with distributional consequences rather than neutral specifications. Buthe and "
  "Mattli (2011) show how first-mover advantage, institutional access and technical capacity determine "
  "whose preferences become embedded in global rules: a body that issues a technically mature proposal "
  "before competitors can mount an alternative sets the text around which others negotiate. Yates and "
  "Murphy (2019) trace the same pattern across a century of voluntary consensus standard-setting, in "
  "which widely adopted committee standards became de facto global norms without formal imposition.")
body("The post-quantum case reproduces that pattern. When the United States National Institute of "
  "Standards and Technology published its first three post-quantum standards in 2024 (NIST, 2024), a "
  "contested question about which algorithms count as quantum-safe became a settled technical fact that "
  "other jurisdictions either adopt or must interoperate with. Standards perform distinct economic "
  "functions, compatibility, quality, variety reduction, and information (Blind, 2013), and the "
  "post-quantum standards act chiefly as compatibility and quality standards. The comparative vocabulary "
  "the Atlas uses, standard-setter, standard-contextualiser, and standard-taker, follows from the power "
  "asymmetry that Buthe and Mattli (2011) describe rather than from a single published typology: a state "
  "can set the baseline, adopt it as given, or occupy an intermediate position of contextualisation. "
  "Bradford's (2020) account of the "
  "Brussels Effect, in which European regulation diffuses globally through market access, is the "
  "instructive contrast: in cryptography the direction reverses, because the standards originate in "
  "Washington and travel through the global deployment of United States technology firms, while European "
  "institutions contextualise rather than set.")
body("Sitting beneath the standards reading is the governance-of-expectations literature (Borup et al., "
  "2006). It explains how shared expectations about an uncertain future coordinate action when settled "
  "knowledge is unavailable (Konrad, 2006), and how that coordination can harden into binding commitment "
  "while remaining revisable (Budde and Konrad, 2019). Because the quantum threat's timing is genuinely "
  "unknowable, the governance of the transition does not consolidate in the way the standards literature "
  "alone would predict; it stays provisional. Carlsson's (2006) work on the internationalisation of "
  "innovation systems frames the cross-border circulation that results. Together these strands give the "
  "Atlas its categories and its central claim: that a leaderless, systemic migration is being coordinated "
  "through expectations, standards and regulation rather than through any global authority.")

# ---- 3 ----
h1("3. The two-axis classification")
body("Every recorded country is described on two dimensions. The coordination posture answers which "
  "collective transition the country belongs to, and takes one of four values: the EU coordinated "
  "roadmap, the NIST-led bloc, a sovereign bloc, or engaged but unaligned. The standards role answers the "
  "country's relationship to the algorithms, and takes one of four values: standard-setter, "
  "standard-contextualiser, standard-taker, or sovereign developer. The posture is the map colour; the "
  "role is a profile badge.")
body("The two are kept separate because one label cannot carry both without losing information. An EU "
  "member coordinates through the European roadmap, yet the algorithms it deploys are the NIST standards, "
  "since no European alternative exists. Describing it with a single camp forces a choice between the club "
  "it belongs to and the technology it runs. Figure 1 shows the pattern across the 54 countries. The "
  "sovereign bloc and the engaged group each sit on a single role, but the NIST and EU blocs split across "
  "roles: the United States is the lone setter inside the NIST bloc, France and Germany contextualise "
  "within the EU roadmap, and the majority of EU members are takers. That split is precisely what a "
  "one-colour scheme would hide.")
figure("fig1_matrix.png", "Figure 1. The two-axis classification: coordination bloc by standards role (54 countries).", 16.0)
body("Each value carries a plain meaning. A setter writes the baseline others adopt. A contextualiser "
  "accepts that baseline but adds national requirements, such as mandatory hybridisation, larger key "
  "sizes, or independent certification. A taker adopts the baseline broadly as published. A sovereign "
  "developer designs and mandates its own algorithms. On the coordination side, the EU roadmap is a "
  "coordination commitment under a shared timeline, the NIST-led bloc works to the United States cadence, "
  "the sovereign bloc runs an independent track, and the engaged-but-unaligned group is active on the "
  "transition without having committed to a bloc, a standard, or a timeline.")

# ---- 4 ----
h1("4. Method")
body("The Atlas is built from institutional sources only: government bodies, national agencies and "
  "standards organisations. Each country's classification is traceable to the documents on its profile, "
  "and a country with no institutional source recedes as no-data rather than being assigned a camp by "
  "inference. Classification follows a single rubric applied in a fixed order. The sovereign test comes "
  "first and is strict: a state is sovereign only where it develops or mandates its own national "
  "post-quantum algorithm, not where it issues a national policy or programme that adopts the "
  "international baseline. EU membership in the coordinated roadmap is next, then adoption of the NIST "
  "suite, then a mixed national-plus-NIST position. A state that is active but has named no standard is "
  "recorded as engaged but unaligned rather than defaulted into the largest bloc.")
body("Two further fields make the picture usable for policy. A confidence grade (high, medium, low) "
  "drives the opacity of each country on the map, so that soft calls read as soft, and a verification "
  "status records whether a claim has been checked against its source. A regulatory layer then records, "
  "for each country, the main governing instrument, its legal status, and the plain obligation it places "
  "on regulated entities. The legal status separates a binding instrument from soft law and from agency "
  "guidance, a distinction without which a recommendation and a statute would read alike. Figure 4 shows "
  "the distribution: most recorded countries carry at least one binding instrument, a single case rests "
  "on soft law alone, and nine have no governing instrument on record.")
figure("fig4_legal.png", "Figure 4. Legal status of the governing instrument across the recorded countries.", 13.5)
h2("AI assistance and continuous updating")
body("The corpus was gathered and provisionally classified with AI assistance under a documented "
  "protocol. Source discovery, extraction and a first-pass classification were produced with model "
  "support; each classification was then checked against its institutional source in an adversarial "
  "verification pass that flagged misattributions and low-confidence defaults for correction. Analytical "
  "authority rests with the researcher: contested codings are decided by hand, and stakeholder material "
  "from the CEPS Task Force is used only as anonymised aggregate findings, with no participant "
  "identification. The result is an audit trail in which every published claim can be traced to a source "
  "and a decision.")
body("The Atlas is a static site with a real geographic map. Country data is held in a version-controlled "
  "store mirrored to a working database, and the site is rebuilt as new institutional documents appear, "
  "so it reflects the current public record rather than a single snapshot. Because the transition is "
  "unfinished, this design matters: the classification is provisional by construction, and the instrument "
  "is meant to move as the evidence does.")

# ---- 5 ----
h1("5. What the map shows")
body("Read by coordination bloc, the 54 recorded countries fall into four groups of very different size "
  "(Figure 2). The EU coordinated roadmap is the largest at 27 members, the NIST-led bloc follows at 15, "
  "nine countries are engaged but unaligned, and three form the sovereign bloc. Read by standards role "
  "(Figure 3), the asymmetry is sharper still: one country sets the baseline, six contextualise it, 35 "
  "take it, three build their own, and nine have not yet declared a role.")
figure("fig2_blocs.png", "Figure 2. Countries by coordination bloc.", 14.5)
figure("fig3_roles.png", "Figure 3. Countries by standards role: one setter, many takers.", 14.5)
body("The United States is the sole standard-setter. To have transitioned as a setter is to write the "
  "algorithms the rest of the world references and to bind one's own agencies to them, while the private "
  "sector follows through procurement and global product deployment rather than a single mandate. The EU "
  "coordinated roadmap is a coordination bloc, not a standards bloc: France and Germany contextualise the "
  "baseline through national certification and technical guidance, while the majority of members are "
  "takers. To have transitioned as an EU taker is to be bound by the Union's general obligations under "
  "NIS2 and DORA, to migrate on the shared 2030 and 2035 timeline, and to run the NIST algorithms with "
  "little national overlay.")
body("The NIST-led bloc gathers the adopters outside the European coordination, including the United "
  "Kingdom and Canada. Four of its members contextualise rather than simply take: India, Japan and "
  "Singapore add national certification and sectoral guidance, and South Korea runs its own national "
  "competition that selects algorithms cognate with the NIST suite while keeping interoperability with "
  "it. The sovereign bloc is the small group that genuinely forks the baseline. China, Russia and Vietnam "
  "develop and mandate their own algorithms, and to have transitioned this way is to accept the "
  "interoperability and market-access costs of a parallel standards track in exchange for cryptographic "
  "independence. The nine engaged-but-unaligned states are active on the transition but have named no "
  "standard or timeline; recording them as undecided keeps the size of each committed bloc honest rather "
  "than inflating the NIST orbit with countries that have not chosen.")

# ---- 6 ----
h1("6. Why the stances matter")
body("The theoretical interest is that the case extends governance-of-expectations and standards-"
  "governance accounts into a transnational, multi-level setting they have not previously addressed. "
  "Expectations generated within one national system, the United States, acquire binding force elsewhere "
  "through market mechanisms and institutional referencing rather than formal agreement (Buthe and "
  "Mattli, 2011). The internationalisation of innovation systems describes that cross-border circulation "
  "(Carlsson, 2006). Reading each country's posture and role is how that circulation becomes "
  "visible as data rather than assertion.")
body("The methodological interest is that separating coordination from role avoids a bloc illusion. "
  "Without the split, the EU reads as a rival standards camp, when on the algorithm axis it runs the same "
  "baseline as the NIST bloc under a regulatory wrapper. The sources-only rule and the confidence grade "
  "keep the claims defensible, and the regulatory layer prevents a soft recommendation from being read as "
  "a binding rule. These choices are what allow the map to say only what the evidence supports.")
body("The geopolitical interest is the most direct. Whoever sets the baseline holds structural power over "
  "the terms on which everyone else secures their systems (Buthe and Mattli, 2011), which is why the "
  "single-setter finding is the central fact of the field rather than a detail. The sovereign fork raises "
  "the prospect of competing spheres of cryptographic interoperability, where divergence in the "
  "underlying standards becomes a barrier to market access. Trade agreements have themselves extended "
  "into regulatory territory in ways that distribute advantage to the already established (Rodrik, 2018). "
  "The European "
  "position as contextualiser, and the contested middle of states that have not yet chosen, are where the "
  "shape of the next few years will be decided.")

# ---- 7 ----
h1("7. Conclusion")
body("The QSC Atlas is a living instrument for a transition that has not finished and may not finish on "
  "any predictable schedule, since no cryptographically relevant quantum computer yet exists. Its "
  "contribution is to make the structure of a leaderless migration legible: which collective each country "
  "coordinates with, whose algorithms it runs, what binds it, and how confident the reading is. The "
  "classification recorded here is provisional, and the instrument is designed to move as institutional "
  "sources appear. The appendix sets out the current position for all 54 recorded countries.")

# ---- appendix (landscape) ----
ap = doc.add_section(WD_SECTION.NEW_PAGE)
ap.orientation = WD_ORIENT.LANDSCAPE
ap.page_width = Cm(29.7); ap.page_height = Cm(21.0)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(ap, m, Cm(1.6))
doc.add_heading("Appendix A. Country classification (54 recorded countries)", level=1)
intro = doc.add_paragraph()
ir = intro.add_run("Coordination bloc and standards role from the QSC Atlas; main regulation, legal status "
  "and obligation from its regulatory layer. Source for all entries: institutional documents on each "
  "country profile.")
ir.italic = True; ir.font.size = Pt(8.5); ir.font.color.rgb = MUTED

rows = json.load(open(os.path.join(ASSETS, "appendix.json")))
cols = [("Country", 3.0), ("Bloc", 3.0), ("Role", 3.6), ("Legal", 2.4), ("Main regulation", 7.0), ("Obligation", 7.1)]
table = doc.add_table(rows=1, cols=len(cols))
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
tbl = table._tbl
layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed'); tbl.tblPr.append(layout)

def set_cell(cell, text, w_cm, bold=False, size=8.0, color=INK):
    cell.width = Cm(w_cm)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(1); p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text); r.bold = bold; r.font.name = "Arial"; r.font.size = Pt(size); r.font.color.rgb = color
    return cell

hdr = table.rows[0].cells
for i, (name, w) in enumerate(cols):
    set_cell(hdr[i], name, w, bold=True, size=8.5)
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), 'EFECE6')
    hdr[i]._tc.get_or_add_tcPr().append(shd)

for row in rows:
    cells = table.add_row().cells
    vals = [row["country"], row["posture"], row["role"] or "Not yet declared", row["legal"] or "-",
            row["reg"] or "None on record", (row["obligation"] + " ") if row["obligation"] else "-"]
    for i, (name, w) in enumerate(cols):
        set_cell(cells[i], vals[i], w)

doc.add_paragraph()

# ---- references (back to portrait) ----
rp = doc.add_section(WD_SECTION.NEW_PAGE)
rp.orientation = WD_ORIENT.PORTRAIT
rp.page_width = Cm(21.0); rp.page_height = Cm(29.7)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(rp, m, Cm(2.4))
doc.add_heading("References", level=1)
refs = [
 "Blind, K. (2013). The impact of standardization and standards on innovation (Nesta Working Paper No. 13/15). Nesta.",
 "Borup, M., Brown, N., Konrad, K., & van Lente, H. (2006). The sociology of expectations in science and technology. Technology Analysis & Strategic Management, 18(3-4), 285-298.",
 "Bradford, A. (2020). The Brussels effect: How the European Union rules the world. Oxford University Press.",
 "Budde, B., & Konrad, K. (2019). Tentative governing of fuel cell innovation in a dynamic network of expectations. Research Policy, 48(5), 1098-1112.",
 "Buthe, T., & Mattli, W. (2011). The new global rulers: The privatization of regulation in the world economy. Princeton University Press.",
 "Carlsson, B. (2006). Internationalization of innovation systems: A survey of the literature. Research Policy, 35(1), 56-67.",
 "European Commission. (2024). Commission Recommendation (EU) 2024/1101 on a coordinated implementation roadmap for the transition to post-quantum cryptography.",
 "European Parliament and Council. (2022a). Directive (EU) 2022/2555 (NIS2).",
 "European Parliament and Council. (2022b). Regulation (EU) 2022/2554 (DORA).",
 "Konrad, K. (2006). The social dynamics of expectations. Technology Analysis & Strategic Management, 18(3-4), 429-444.",
 "NIS Cooperation Group. (2025). Coordinated implementation roadmap for the transition to post-quantum cryptography.",
 "NIST. (2024). FIPS 203, FIPS 204 and FIPS 205: Post-quantum cryptography standards. National Institute of Standards and Technology.",
 "Rodrik, D. (2018). What do trade agreements really do? Journal of Economic Perspectives, 32(2), 73-90.",
 "Star, S. L. (1999). The ethnography of infrastructure. American Behavioral Scientist, 43(3), 377-391.",
 "Yates, J., & Murphy, C. N. (2019). Engineering rules: Global standard setting since 1880. Johns Hopkins University Press.",
]
for ref in refs:
    p = doc.add_paragraph(); r = p.add_run(ref); r.font.size = Pt(9.5)
    p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_after = Pt(5)

doc.save(OUT)
print("saved:", OUT)
print("paragraphs:", len(doc.paragraphs), "| appendix rows:", len(rows))
