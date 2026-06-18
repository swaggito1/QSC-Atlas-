#!/usr/bin/env python3
# Builds the QSC Atlas expert engagement and validation plan (.docx).
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/sessions/gracious-kind-wozniak/mnt/qsc-atlas/docs/QSC_Atlas_Expert_Engagement_Plan.docx"
INK = RGBColor(0x1A, 0x1A, 0x1A); MUTED = RGBColor(0x5C, 0x5C, 0x5C)

doc = Document()
n = doc.styles["Normal"]; n.font.name = "Georgia"; n.font.size = Pt(10.5); n.font.color.rgb = INK
n.paragraph_format.space_after = Pt(7); n.paragraph_format.line_spacing = 1.22
for hid, sz in (("Heading 1", 14), ("Heading 2", 11.5)):
    st = doc.styles[hid]; st.font.name = "Arial"; st.font.size = Pt(sz); st.font.bold = True; st.font.color.rgb = INK
    st.paragraph_format.space_before = Pt(12); st.paragraph_format.space_after = Pt(5)

sec = doc.sections[0]
sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
for m in ("top_margin","bottom_margin","left_margin","right_margin"): setattr(sec, m, Cm(2.3))

def body(text, after=7):
    p = doc.add_paragraph(); p.add_run(text); p.paragraph_format.space_after = Pt(after); return p
def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)

def table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
    lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'),'fixed'); t._tbl.tblPr.append(lay)
    def cell(c, txt, w, bold=False, fill=None, size=9.0):
        c.width = Cm(w); p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1); p.paragraph_format.line_spacing = 1.0
        r = p.add_run(txt); r.bold = bold; r.font.name = "Arial"; r.font.size = Pt(size); r.font.color.rgb = INK
        c.margins = None
        if fill:
            shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),fill); c._tc.get_or_add_tcPr().append(shd)
    for i,hh in enumerate(headers): cell(t.rows[0].cells[i], hh, widths[i], bold=True, fill="EFECE6", size=9.0)
    for row in rows:
        cs = t.add_row().cells
        for i,val in enumerate(row): cell(cs[i], val, widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

# ---- title ----
p = doc.add_paragraph(); r = p.add_run("QSC Atlas: expert engagement and validation plan"); r.font.name="Georgia"; r.font.size=Pt(18); r.bold=True
p.paragraph_format.space_after = Pt(2)
p = doc.add_paragraph(); r = p.add_run("Stakeholder review across summer 2026, with a validated output by end September 2026"); r.font.name="Arial"; r.font.size=Pt(11.5); r.font.color.rgb=MUTED
p.paragraph_format.space_after = Pt(2)
p = doc.add_paragraph(); r = p.add_run("Draft plan, June 2026. Companion to the QSC Atlas (qsc-atlas.vercel.app). [Author and affiliation]"); r.font.name="Arial"; r.font.size=Pt(9); r.font.color.rgb=MUTED
p.paragraph_format.space_after = Pt(12)

# ---- 1 ----
h1("1. Purpose and objective")
body("The QSC Atlas and its companion report are complete as a first edition. This plan turns that "
  "edition into a collaborative, expert-validated instrument. Over the summer, fifteen senior figures "
  "from industry, the public sector and international organisations will review the Atlas and say whether "
  "it is relevant and what is wrong, and propose how to improve it. Their input is incorporated through a "
  "documented process, so the second edition carries both a sharper classification and the credibility of "
  "named expert review.")
body("The objective is to tailor the Atlas as closely as possible to how practitioners and policymakers "
  "actually read the quantum-safe transition: to test the accuracy of each classification, the soundness "
  "of the two-axis method, and the usefulness of the regulatory layer, and to correct what the evidence "
  "and the experts show to be wrong. A secondary objective is to build the relationships and the standing "
  "that a validated, expert-backed instrument needs to grow into an advisory practice.")

# ---- 2 ----
h1("2. What the experts review")
body("Reviewers see a short pack: the live Atlas, the companion report, and a one-page summary of the "
  "method. They are asked to assess five things.")
table(["Object of review", "What we want tested"],
 [["The two-axis classification","Whether coordination posture and standards role are the right axes, clearly defined and correctly separated."],
  ["Per-country classification","Whether each country's coordination bloc and standards role are accurate, and which are misattributed."],
  ["The regulatory layer","Whether the main regulation, legal status and obligation recorded for each country are correct and current."],
  ["The methodology","Whether the sources-only rule, the rubric and decision order, and the confidence grading are defensible."],
  ["Framing and theory","Whether the report's standards-governance framing is accurate and useful for the intended audiences."]],
 [4.6, 11.7])

# ---- 3 ----
h1("3. The panel: fifteen experts")
body("Five reviewers from each of three constituencies, chosen for seniority, independence, and spread "
  "across the coordination blocs (NIST-led, EU, sovereign-adjacent, and engaged-but-unaligned). Targets "
  "are described by role and organisation type; specific individuals are confirmed in Phase 0.")
table(["Constituency", "N", "Target profiles (by role and organisation type)"],
 [["Industry","5","Cryptography or PKI vendor; cloud or platform security lead; financial-institution crypto or CISO lead; telecoms or GSMA-track engineer; post-quantum start-up or hardware security module maker."],
  ["Public sector","5","National cyber agency cryptography leads (two, from different blocs); a financial-sector regulator; a defence or critical-infrastructure body; a representative from a non-Western or engaged-but-unaligned state."],
  ["International and standards","5","ENISA; the ETSI quantum-safe group; ISO/IEC JTC 1 SC 27; a multilateral financial body (for example the BIS or the G7 Cyber Expert Group); a regional or OECD-level convenor or a post-quantum coalition."]],
 [3.7, 0.9, 11.7])
body("Selection balances the blocs so the review is not dominated by one camp, and mixes standards "
  "leads, policy leads and operators so relevance is tested for more than one kind of reader. Conflicts "
  "of interest are declared at recruitment, and reviewers choose whether to be named or to remain "
  "anonymous in the published acknowledgements.")

# ---- 4 ----
h1("4. Engagement model")
body("The review follows a light, two-round expert-elicitation design (a modified Delphi). The first "
  "round is individual and asynchronous, which respects senior schedules and avoids one voice anchoring "
  "the rest. The second round convenes the panel to resolve the points where reviewers disagree, so the "
  "outcome reflects considered consensus rather than a first reaction. The whole ask is kept to about one "
  "hour per reviewer in round one, with an optional thirty-minute interview for those who prefer to talk.")
body("Reviewers are offered acknowledgement as named expert reviewers, early access to the second "
  "edition, and a contribution to a short published validation note. The roundtable runs under the "
  "Chatham House rule, so positions can be discussed candidly and reported without attribution.")

# ---- 5 ----
h1("5. The feedback instrument: two questions and a route to improvement")
body("Every reviewer answers the same two questions, with a structured path from a criticism to an "
  "actionable, sourced suggestion.")
table(["Question", "What is asked", "How the reviewer answers"],
 [["Q1. Is it relevant?","Whether the Atlas is useful, for what purpose, and for whom.","A one-to-five rating with a one-line statement of the use case it serves or fails to serve."],
  ["Q2. What is wrong?","What is inaccurate, missing or misleading, across the five objects of review.","Structured flags: for classification, the specific country, its current camp and the suggested camp, with a reason; for the method and the regulatory layer, the specific point contested."],
  ["Route to improvement","How to make it more accurate and more authoritative.","For each flag, a 'suggested change plus supporting source plus confidence' field, and one overall prompt: what single change would make this authoritative for your organisation."]],
 [3.3, 5.3, 7.7])
body("Reviewers choose the channel that suits them: a structured online form, inline comments on a "
  "shared review copy, a 'suggest a correction' path on the live Atlas keyed to each country, or a "
  "recorded interview that the team transcribes. All channels feed a single disposition log, so every "
  "comment is captured with its source and routed to a decision. Capturing the source and a confidence "
  "level on each suggestion lets the change flow straight into the Atlas verification process.")

# ---- 6 ----
h1("6. Timeline")
body("Engagement runs from mid-June to 15 September, the hard deadline for all reviewer input, with the "
  "validated edition delivered by the end of September.")
table(["Phase", "Dates (2026)", "Activities", "Output"],
 [["0. Set-up and recruitment","17 Jun to 4 Jul","Finalise the review pack and the form; confirm the fifteen experts; consent and conflict declarations; brief the panel.","Confirmed panel; review pack; live feedback instrument."],
  ["1. Round 1, individual review","7 to 25 Jul","Distribute the pack and form; run optional interviews; send reminders.","Fifteen completed reviews."],
  ["2. Synthesis and draft revision","28 Jul to 15 Aug","Code the feedback; build the disposition log; draft reclassifications and method revisions; mark contested points.","'What we heard' synthesis; draft second edition."],
  ["3. Round 2, convening and resolution","18 Aug to 12 Sep","Roundtable under the Chatham House rule and targeted one-to-ones on contested points; collect final corrections.","Resolved positions; final corrections. Input closes 15 Sep."],
  ["4. Finalisation and delivery","15 to 30 Sep","Apply changes; re-verify against sources; finalise the outputs.","Validated edition delivered end September."]],
 [3.3, 2.6, 6.4, 4.0])

# ---- 7 ----
h1("7. Outputs, end September")
body("The validated edition has five parts: the revised Atlas, with reclassifications applied and "
  "confidence grades updated; the companion report, with a new section on the expert validation; a short "
  "validation and methodology note that documents the process and shows how each strand of feedback was "
  "used; the named (or anonymised) panel of expert reviewers, as a credibility layer; and the disposition "
  "log, which records every comment and the decision taken on it. The log matters: it is what lets the "
  "project say the Atlas was reviewed by named experts and show exactly what changed as a result.")

# ---- 8 ----
h1("8. Governance, ethics and authority")
body("Participation is by informed consent, with each reviewer choosing named or anonymous attribution. "
  "Conflicts of interest are declared at recruitment and recorded. The roundtable runs under the Chatham "
  "House rule. Any material drawn from closed stakeholder settings is used only as anonymised aggregate, "
  "consistent with the project's existing confidentiality commitments. Analytical authority stays with "
  "the project lead: experts advise, and the lead decides and documents each decision in the disposition "
  "log. Reviewers are advising on a research instrument, not signing off a compliance product.")

# ---- 9 ----
h1("9. Contributors needed beyond the report")
body("To make the Atlas high-level and to grow its commercial capacity, the project needs contributors "
  "in two groups: those who add authority and keep the instrument accurate, and those who turn it into a "
  "service. The table marks which roles are needed first.")
h2("Credibility and quality")
table(["Role", "What they add", "Priority"],
 [["Academic or methodological co-lead, or advisory chair","A recognised authority in standards governance or cryptography who co-signs the method and lends standing.","First"],
  ["Expert advisory board (three to five)","Ongoing governance and credibility beyond the one-off review; a standing source of correction.","First"],
  ["Data and research analyst","Maintains the corpus, runs the country research, and keeps the classifications current.","First"],
  ["Software and data engineer","Maintains the site and the update pipeline; builds the 'suggest a correction' tool and automates the rebuilds.","First"],
  ["Policy and sector analyst","Turns the data into sector briefs (finance, critical infrastructure) where the obligations bite.","Second"],
  ["Editorial support","Keeps the report, briefs and analysis channel to a regular cadence.","Second"]],
 [4.7, 9.1, 2.5])
h2("Reach and commercial capacity")
table(["Role or partner", "What they add", "Priority"],
 [["Commercial or business-development lead","Productises the Atlas: country briefs, a monitoring subscription, bespoke deep-dives; owns partnerships and sales.","First"],
  ["Partnerships and communications lead","Institutional anchors (a host such as CEPS or a university, and standards bodies) and dissemination.","First"],
  ["Legal and compliance adviser","Keeps any advisory output sound, and handles data, IP and disclaimers.","Second"],
  ["Host institution and funding or sponsor partner","An institutional home that confers authority, and resourcing for the engagement and the platform.","First"]],
 [4.7, 9.1, 2.5])
body("The minimum viable team to reach a credible, sellable second edition is small: a methodological "
  "co-lead or advisory chair, a data analyst, an engineer, and a commercial lead, anchored to a host "
  "institution. The advisory board and the sector and editorial roles can follow once the validated "
  "edition is in hand and the first advisory engagements are running.")

doc.save(OUT)
print("saved:", OUT, "| paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables))
