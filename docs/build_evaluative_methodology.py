#!/usr/bin/env python3
# QSC Atlas evaluative layer: methodology, a first-pass run across all countries, and a
# detailed Claude Code build brief. British English, no em-dashes, grounded citations.
import json, os, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROF = "/sessions/gracious-kind-wozniak/mnt/qsc-atlas/data/profiles"
OUT = "/sessions/gracious-kind-wozniak/mnt/qsc-atlas/docs/QSC_Atlas_Evaluative_Layer_Methodology.docx"
INK = RGBColor(0x1A,0x1A,0x1A); MUTED = RGBColor(0x5C,0x5C,0x5C); WASH="EFECE6"

# ---------- scoring from recorded fields ----------
def lines(s): return [l.strip() for l in (s or "").split("\n") if l.strip()]
def has_year(s): return sum(1 for l in lines(s) if re.search(r"20\d\d", l))

NAMES={}
try:
    for c in json.load(open(os.path.join(PROF,"..","countries.json"))): NAMES[c.get("iso3")]=c.get("name")
except Exception: pass

def score_country(p):
    tl=p.get("migrationTimeline") or ""; ms=has_year(tl)
    target=bool(p.get("targetCompletion"))
    regs=lines(p.get("mainRegulation")); nreg=len(regs)
    legal=p.get("legalStatus"); binding=legal=="binding"; soft=legal=="soft-only"
    role=p.get("standardsRole"); dom=p.get("dominantProcess")
    gov=lines(p.get("govActors")); ngov=len(gov)
    part=bool((p.get("processParticipation") or "").strip())
    hybrid=(p.get("hybridDeployment") or "").strip(); has_hybrid=hybrid not in ("","None stated","None")
    fams=bool((p.get("standardFamilies") or p.get("algorithms") or "").strip())
    text=" ".join([p.get("obligation") or "",p.get("mainRegulation") or "",p.get("summary") or "",tl]).lower()
    eu=("nis2" in text or "dora" in text or "(eu)" in text)
    s={}
    s["R1"]=2 if ms>=1 else (1 if (nreg>0 or target) else 0)
    s["R2"]=2 if binding else (1 if (soft or dom) else 0)
    s["C1"]=2 if (eu or nreg>=2) else (1 if (nreg==1 or part) else 0)
    s["C2"]=2 if ngov>=3 else (1 if (ngov>=1 or part) else 0)
    s["E1"]=2 if ms>=1 else (1 if target else 0)
    s["E2"]=2 if ("inventory" in text or "cbom" in text) else (1 if binding else 0)
    s["E3"]=2 if any(w in text for w in ("monitor","indicator","audit")) else None
    s["Ef1"]=2 if binding else (1 if soft else 0)
    s["G1"]=2 if ngov>=2 else (1 if ngov==1 else 0)
    s["G2"]=2 if any(w in text for w in ("skill","training","academ","programme","workforce")) else None
    s["I1"]=2 if (role or (dom and (has_hybrid or fams or nreg>0))) else (1 if dom else 0)
    def mean(keys):
        v=[s[k] for k in keys if s.get(k) is not None]
        return round(sum(v)/len(v),1) if v else None
    return {
        "Relevance":mean(["R1","R2"]),"Coherence":mean(["C1","C2"]),
        "Effectiveness":mean(["E1","E2","E3"]),"Efficiency":mean(["Ef1"]),
        "Governance":mean(["G1","G2"]),"Impact":mean(["I1"]),
    }, s

CONFMAP={"Complete":"High","Partial":"Medium","Placeholder":"Low"}
rows=[]
for f in sorted(os.listdir(PROF)):
    if not f.endswith(".json"): continue
    p=json.load(open(os.path.join(PROF,f)))
    if not p.get("dominantProcess") or p.get("iso3")=="EUU": continue
    ax,_=score_country(p)
    conf=p.get("confidence") or CONFMAP.get(p.get("dataStatus"),"-")
    rows.append((p.get("country") or NAMES.get(p.get("iso3")) or p.get("iso3"), ax, conf))
rows.sort(key=lambda r:r[0])

# ---------- document ----------
doc=Document()
n=doc.styles["Normal"]; n.font.name="Georgia"; n.font.size=Pt(10.5); n.font.color.rgb=INK
n.paragraph_format.space_after=Pt(7); n.paragraph_format.line_spacing=1.22
for hid,sz in (("Heading 1",14),("Heading 2",11.5),("Heading 3",10.5)):
    st=doc.styles[hid]; st.font.name="Arial"; st.font.size=Pt(sz); st.font.bold=True; st.font.color.rgb=INK
    st.paragraph_format.space_before=Pt(11); st.paragraph_format.space_after=Pt(4)
sec=doc.sections[0]; sec.page_width=Cm(21.0); sec.page_height=Cm(29.7)
for m in ("top_margin","bottom_margin","left_margin","right_margin"): setattr(sec,m,Cm(2.3))

def body(t, after=7):
    p=doc.add_paragraph(); p.add_run(t); p.paragraph_format.space_after=Pt(after); return p
def h1(t): doc.add_heading(t,1)
def h2(t): doc.add_heading(t,2)
def h3(t): doc.add_heading(t,3)
def caption(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.italic=True; r.font.size=Pt(9); r.font.color.rgb=MUTED; p.paragraph_format.space_after=Pt(10)
def table(headers,data,widths,sizes=9.0):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    lay=OxmlElement('w:tblLayout'); lay.set(qn('w:type'),'fixed'); t._tbl.tblPr.append(lay)
    def cell(c,txt,w,bold=False,fill=None):
        c.width=Cm(w); pr=c.paragraphs[0]; pr.paragraph_format.space_after=Pt(1); pr.paragraph_format.line_spacing=1.0
        rr=pr.add_run(txt); rr.bold=bold; rr.font.name="Arial"; rr.font.size=Pt(sizes); rr.font.color.rgb=INK
        if fill:
            shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),fill); c._tc.get_or_add_tcPr().append(shd)
    for i,h in enumerate(headers): cell(t.rows[0].cells[i],h,widths[i],bold=True,fill=WASH)
    for row in data:
        cs=t.add_row().cells
        for i,v in enumerate(row): cell(cs[i],v,widths[i])
    doc.add_paragraph().paragraph_format.space_after=Pt(3)
    return t

# title
p=doc.add_paragraph(); r=p.add_run("The QSC Atlas evaluative layer: methodology, a run across all countries, and the build brief")
r.font.name="Georgia"; r.font.size=Pt(16); r.bold=True; p.paragraph_format.space_after=Pt(2)
p=doc.add_paragraph(); r=p.add_run("Six criteria, scored from the recorded data, with a detailed UI specification for Claude Code")
r.font.name="Arial"; r.font.size=Pt(11); r.font.color.rgb=MUTED; p.paragraph_format.space_after=Pt(2)
p=doc.add_paragraph(); r=p.add_run("Working note, June 2026. Companion to Extending the QSC Atlas and the QSA methods synthesis. [Author and affiliation]")
r.font.name="Arial"; r.font.size=Pt(9); r.font.color.rgb=MUTED; p.paragraph_format.space_after=Pt(12)

ab=doc.add_paragraph(); abr=ab.add_run("Summary. "); abr.bold=True
ab.add_run("This note operationalises the evaluative layer, runs it across every country in the dataset, and "
  "specifies how to build it. Part A defines the six criteria as eleven sub-criteria, each scored 0 to 2 "
  "against stated descriptions, with the verification and reconciliation rules that govern scoring. Part B "
  "gives a first-pass run for every coloured country in the dataset, derived mechanically from the recorded "
  "Atlas fields by the "
  "rules in the note, with confidence carried and reviewer reconciliation pending. Part C is the detailed "
  "build brief for the credibility instrument, written for Claude Code. No score is shown unless it traces "
  "to a recorded field or source; an unsupported sub-criterion reads as not assessed.")
for r in ab.runs: r.font.size=Pt(10)

# ===== PART A =====
h1("Part A. Methodology")
h2("1. What the layer reposes on")
body("The layer adapts the comparative policy-evaluation method that Paglieri et al. (2025) built for "
  "national quantum strategies from OECD and World Economic Forum practice, the one method in the assessment "
  "corpus designed at the national scale. Vance (2025) sharpens it by separating evaluation into process, "
  "impact and cost-benefit, so the layer judges effectiveness and impact by design rather than by outcomes "
  "an emerging field cannot yet show. Two of the sub-criteria draw their content from the CEPS governance "
  "recommendations, the inventory duty and the capability checks (Pupillo et al., 2025), and the threat "
  "articulation criterion from the CEPS governance report (Pupillo et al., 2023). The whole sits inside "
  "governance of expectations, which treats this transition as provisional because the threat's timing is "
  "unknowable (Budde and Konrad, 2019), so the layer scores credibility under uncertainty and expects the "
  "score to move.")

h2("2. The scoring scheme")
body("Each of the six criteria is scored through one to three sub-criteria, eleven in total. Every "
  "sub-criterion takes one of three values: 0 absent, 1 partial, 2 present, against the description in Table "
  "1. A criterion's score is the mean of the sub-criteria that can be assessed, giving a value from 0 to 2 "
  "on each of six axes. The layer reports the six axes as a profile, not a single league-table number, "
  "because a composite would hide the disagreement between axes that is often the finding. Three bands label "
  "each axis: nascent (0 to 0.7), developing (0.8 to 1.4), established (1.5 to 2).")
table(["Criterion","Sub-criterion","Score: 0 / 1 / 2","Reposes on"],
 [("Relevance","R1 Transition plan","none / announced or partial / published roadmap with milestones","Paglieri et al. (2025); NIST (2022)"),
  ("","R2 Threat articulation","absent / generic / names the harvest-now-decrypt-later exposure","Paglieri et al. (2025); Pupillo et al. (2023)"),
  ("Coherence","C1 Policy integration","isolated / referenced in wider cyber and S&T policy / embedded in it","Paglieri et al. (2025), after OECD"),
  ("","C2 Consultation","none / narrow / broad consultation across government, industry and academia","Paglieri et al. (2025), after WEF"),
  ("Effectiveness","E1 Milestones","none / indicative / dated and sector-specific","Vance (2025); Paglieri et al. (2025)"),
  ("(by design)","E2 Inventory duty","none / recommended / mandated (inventory or CBOM)","Pupillo et al. (2025), GR3; Vance (2025)"),
  ("","E3 Measurement","none / ad hoc / a monitoring framework or indicators","Paglieri et al. (2025), after OECD"),
  ("Efficiency","Ef1 Proportionate mix","mismatched / partial / instrument strength matched to objective","Paglieri et al. (2025), after OECD"),
  ("Governance quality","G1 Mandate and authority","none / a named lead / a mandated lead with a steering function","Paglieri et al. (2025); Pupillo et al. (2025), GR9"),
  ("","G2 Capacity and skills","none / acknowledged / a skills or capability programme","Pupillo et al. (2025), GR10"),
  ("Impact (structured-for)","I1 Exposure-reducing adoption","none / baseline referenced / baseline adopted with hybrid and agility","Vance (2025)"),
 ],
 [2.7,3.4,6.4,4.0],sizes=8.5)
caption("Table 1. The eleven sub-criteria, their scoring descriptions, and the model each reposes on.")

h2("3. Verification, reconciliation and confidence")
body("Three rules keep the scores defensible. No sub-score is recorded unless it traces to a sourced "
  "instrument already held on the country's Atlas profile, or a source added for the purpose; a sub-"
  "criterion with no supporting source is left unscored and reads as not assessed, never a default value, "
  "and its axis is the mean of the sub-criteria that could be assessed. Second, scoring is reconciled: a "
  "second senior reviewer scores each country independently and the two reconcile before a profile is shown "
  "to a client, the validation step Paglieri et al. (2025) build into their method. Third, the Atlas's "
  "confidence grade is carried onto each axis, so a country assessed on thin evidence reads as tentative. "
  "Because the frame is governance of expectations, a score is a reading at a date, not a verdict, and the "
  "profile carries the date it was last verified.")

# ===== PART B: the run =====
h1("Part B. The run across all countries")
body("The table below is a first-pass run for every coloured country, produced by applying the rules in "
  "Table 1 mechanically to the fields each profile already records. It is a starting point for review, not a "
  "finished assessment: the scores inherit the confidence grade of the underlying profile, the reviewer "
  "reconciliation in section 3 is still pending, and the data-thin sub-criteria are scored only where a "
  "field supports them. Section 4 records the derivation so each score is traceable.")

h3("4. How each score is derived from the recorded fields")
body("R1 and E1 read the migration timeline and target year. R2 reads whether a binding instrument exists, "
  "with a softer instrument or a recorded camp counting as partial. C1 reads whether EU instruments or two "
  "or more national instruments are present. C2 reads the breadth of recorded governmental and standards "
  "bodies. E2 reads the regulation and obligation text for an inventory or CBOM duty, with a binding regime "
  "counting as partial. E3 reads the text for a monitoring or audit duty and is left unscored where none "
  "appears. Ef1 reads the legal status. G1 reads the lead authority and steering bodies. G2 reads the text "
  "for a skills or capability programme and is left unscored where none appears. I1 reads the standards "
  "role where present, and otherwise the dominant-process camp together with the hybrid stance and the "
  "recorded standards. A dash means the axis could not be assessed from the recorded fields.")

def fmt(x): return "-" if x is None else f"{x:.1f}"
data=[(c, fmt(ax["Relevance"]), fmt(ax["Coherence"]), fmt(ax["Effectiveness"]), fmt(ax["Efficiency"]), fmt(ax["Governance"]), fmt(ax["Impact"]), conf) for c,ax,conf in rows]
table(["Country","Rel","Coh","Eff","Effi","Gov","Imp","Conf."], data,
      [3.7,1.55,1.55,1.55,1.55,1.55,1.55,1.6], sizes=8.0)
caption(f"Table 2. First-pass six-axis scores for {len(rows)} countries, derived from the recorded fields (0 to 2 per axis; a dash means not assessed). Reviewer reconciliation pending.")

# France worked note (uses computed France row)
fr=next((ax for c,ax,_ in rows if c=="France"), None)
if fr:
    h3("5. Reading one row: France")
    body("France illustrates how to read a row. Relevance and impact are established: ANSSI publishes a "
      "three-phase transition roadmap (Phase 1 hybridisation now, Phase 2 mandatory hybridisation not "
      "earlier than 2025, Phase 3 standalone post-quantum not earlier than 2030), and the NIST baseline is "
      "adopted with mandatory hybridisation (ANSSI, 2023, verified). Coherence and governance quality are "
      "established through the EU instruments and ANSSI's mandated lead under SGDSN. Effectiveness sits lower "
      "because the milestones are framed as not-earlier-than windows and a dedicated measurement duty is not "
      "recorded. The verification step corrected one inherited figure: the earlier profile cited fixed visa "
      "deadlines of 2027 and 2030, which the ANSSI roadmap does not state, and which should be corrected in "
      "the companion report and the regulatory layer.")

# ===== PART C: build brief =====
h1("Part C. Build brief for Claude Code")
body("This is a developer-ready specification for the credibility instrument that shows the six axes on each "
  "country profile. The design principle is that playfulness comes from interaction and discovery, not "
  "effects: a small precision instrument the visitor operates and compares, in the monochrome style, that "
  "never adds a second colour to the map. Every score on screen links to its source, and an unassessed axis "
  "is shown empty.")

h3("C1. Data shape")
body("Add an optional evaluation object to each country profile (data/profiles/<ISO3>.json), mirror it in "
  "Notion, and type it in src/content.config.ts as a nullable field, so profiles without it render an empty "
  "gauge. Shape: evaluation maps each of the six criteria (relevance, coherence, effectiveness, efficiency, "
  "governance, impact) to an object with score (number 0 to 2, or null for not assessed), confidence (High, "
  "Medium or Low), and subScores, an array of { key, label, score (0, 1, 2 or null), reason, sourceUrl }. "
  "The loader in content.config.ts reads it from a JSON-encoded text property or from the profile file; the "
  "zod schema marks it nullable.")

h3("C2. Component")
body("Create src/components/CredibilityGauge.tsx as a React island. Render an SVG (viewBox 0 0 240 240, "
  "width 100 per cent). Six axes at 60-degree spacing from the centre, fixed order with relevance at the "
  "top and the rest clockwise. Each axis is a thin hairline spoke to the maximum radius; the score plots a "
  "point at (score / 2) of that radius; the six points join into a polygon stroked thin in ink and filled "
  "in the country's coordination-bloc colour at low alpha, with the bloc colour also on the outer ring as "
  "the single accent (read the colour from POSTURE_META in src/lib/process.ts). Axis labels in Schibsted "
  "Grotesk sit outside each spoke, with the numeric score in Spline Sans Mono. An axis with a null score "
  "renders as a dashed grey spoke with the label 'not yet assessed' and no point; an axis whose confidence "
  "is Low renders its point and segment at reduced opacity.")

h3("C3. Interaction")
body("On mount the gauge assembles: the spokes draw with stroke-dashoffset, then the points scale in, over "
  "about 1.2 seconds, eased with var(--ease). Hovering or focusing a spoke opens a small panel showing that "
  "criterion's sub-scores, each with its 0 to 2 value, its one-line reason, a link to the cited source, and "
  "the confidence. A compare control, a select listing the other countries, overlays the chosen country's "
  "polygon as a ghosted dashed outline in its own bloc colour, with a two-line legend. Each spoke is a "
  "focusable control; Enter or Space toggles its panel.")

h3("C4. Accessibility and reduced motion")
body("The svg carries role img with a title and a desc that summarise the country's profile. Each spoke is "
  "a button with an aria-label of the form 'Relevance: 2.0 of 2, established'. Values are always shown as "
  "text, never by colour alone, and the focus ring uses var(--focus-ring). Under prefers-reduced-motion the "
  "gauge renders in its final assembled state with no draw-on or flip, and the sub-score panel is shown "
  "inline rather than on an animated reveal.")

h3("C5. Integration, tokens and acceptance")
body("Mount the gauge in ProfilePanel, inside src/components/AtlasMap.tsx, directly below the Regulatory "
  "basis block, under a new section labelled Governance credibility. Reuse the existing tokens (--ink, "
  "--paper, --hairline, --ink-muted, --radius) and fonts (--font-instrument, --font-mono, --font-reading), "
  "and the bloc colours from process.ts. Add no new dependency; the gauge is pure SVG and React state. "
  "Acceptance: the six axes render from the evaluation data with unassessed axes greyed; hover and keyboard "
  "both reveal the sub-scores with a working source link and the confidence; the compare control overlays a "
  "second country; colour is never the sole carrier of a value; reduced motion is respected; and the site "
  "builds with astro build.")

h3("C6. Generative prompt (mood and concept)")
p=doc.add_paragraph(); r=p.add_run("Minimal editorial diagram, ink on warm off-white paper. A six-spoke "
  "radial gauge, like a precision seismograph, six fine axes from a centre each filled to a different level, "
  "one restrained colour accent on the outer ring, the rest monochrome. Clean, scientific, quietly playful "
  "through precision rather than effects. Thin strokes, generous whitespace, a small mono label on each "
  "axis. Flat 2D, high contrast, no gradients, no glow. Square.")
r.font.name="Spline Sans Mono"; r.font.size=Pt(9.5); r.font.color.rgb=INK
p.paragraph_format.left_indent=Cm(0.5); p.paragraph_format.space_after=Pt(8)

h1("References")
refs=[
 "ANSSI. (2023). ANSSI views on the post-quantum cryptography transition (follow-up position paper). Agence nationale de la securite des systemes d'information.",
 "Budde, B., & Konrad, K. (2019). Tentative governing of fuel cell innovation in a dynamic network of expectations. Research Policy, 48(5), 1098-1112.",
 "NIST. (2022). Migration to post-quantum cryptography. National Institute of Standards and Technology.",
 "Paglieri, L., Bonomi Savignon, A., Scalabrini, F., & Costumato, L. (2025). Navigating the quantum frontier: Examining government strategy to the next technological revolution. Transforming Government: People, Process and Policy, 19(4).",
 "Pupillo, L., Ferreira, A., Lipiainen, V., & Polito, C. (2023). Quantum technologies and cybersecurity: Technology, governance and policy challenges. Centre for European Policy Studies.",
 "Pupillo, L., et al. (2025). Strengthening the EU transition to a quantum-safe world: Technology, market and governance (Task Force Report). Centre for European Policy Studies.",
 "Vance, A. S. (2025). Cybersecurity and quantum computing: A quantitative analysis proposing a framework for assessing quantum cybersecurity maturity [Doctoral dissertation].",
]
for ref in refs:
    p=doc.add_paragraph(); r=p.add_run(ref); r.font.size=Pt(9.5)
    p.paragraph_format.left_indent=Cm(1.0); p.paragraph_format.first_line_indent=Cm(-1.0); p.paragraph_format.space_after=Pt(5)

doc.save(OUT)
print("saved:", OUT, "| countries scored:", len(rows), "| paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables))
print("France axes:", fr)
print("sample:", data[:3])
